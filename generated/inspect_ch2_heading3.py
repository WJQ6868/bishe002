# -*- coding: utf-8 -*-
from docx import Document
p = r"D:\bishe\one\tmp_thesis_567.docx"
doc = Document(p)
for i in range(68, 125):
    t = doc.paragraphs[i].text.strip()
    if t:
        print(i, t, doc.paragraphs[i].style.name)
