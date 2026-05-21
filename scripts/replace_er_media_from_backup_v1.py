from __future__ import annotations

from pathlib import Path
from zipfile import ZIP_DEFLATED, ZipFile

from docx import Document
from docx.oxml.ns import qn


def paragraph_text(paragraph) -> str:
    return paragraph.text.strip()


def find_media_targets(doc_path: Path, caption_to_image: dict[str, Path]) -> dict[str, Path]:
    doc = Document(str(doc_path))
    media_map: dict[str, Path] = {}
    for index, paragraph in enumerate(doc.paragraphs):
        caption = paragraph_text(paragraph)
        if caption not in caption_to_image or index == 0:
            continue
        image_paragraph = doc.paragraphs[index - 1]
        blips = image_paragraph._p.xpath(".//*[local-name()='blip']")
        if not blips:
            raise RuntimeError(f"Caption '{caption}' found, but previous paragraph has no image in {doc_path}")
        for blip in blips:
            rel_id = blip.get(qn("r:embed"))
            if not rel_id:
                continue
            image_part = doc.part.related_parts[rel_id]
            media_path = image_part.partname.lstrip("/")
            media_map[media_path] = caption_to_image[caption]
    return media_map


def rewrite_docx_with_media(source_doc: Path, dest_doc: Path, media_map: dict[str, Path]) -> None:
    replacements = {name: image_path.read_bytes() for name, image_path in media_map.items()}
    with ZipFile(source_doc, "r") as src, ZipFile(dest_doc, "w", ZIP_DEFLATED) as dst:
        for item in src.infolist():
            data = replacements.get(item.filename, src.read(item.filename))
            dst.writestr(item, data)


def replace_from_backup(source_doc: Path, dest_doc: Path, caption_to_image: dict[str, Path]) -> dict[str, str]:
    media_map = find_media_targets(source_doc, caption_to_image)
    rewrite_docx_with_media(source_doc, dest_doc, media_map)
    return {name: path.name for name, path in media_map.items()}


def main() -> None:
    root = Path(r"D:\bishe\one")
    asset_dir = root / "generated" / "thesis_assets"
    desktop = Path.home() / "Desktop"

    role_image = asset_dir / "fig_2_1_role_demand_er_v2.png"
    ai_image = asset_dir / "fig_3_1_ai_er_compact_v2.png"

    backup_567 = max(desktop.glob("2405273202-王佳齐-AI赋能的高校教务系统(567)-backup-before-er-enlarge-*.docx"))
    backup_v14 = max(desktop.glob("2405273202-王佳齐-AI赋能的高校教务系统(567)-chapter2-updated-v14-backup-before-er-enlarge-*.docx"))

    target_567 = desktop / "2405273202-王佳齐-AI赋能的高校教务系统(567).docx"
    target_v14 = desktop / "2405273202-王佳齐-AI赋能的高校教务系统(567)-chapter2-updated-v14.docx"

    result_567 = replace_from_backup(
        backup_567,
        target_567,
        {
            "图3-1 AI核心数据ER图": ai_image,
        },
    )
    result_v14 = replace_from_backup(
        backup_v14,
        target_v14,
        {
            "图二-1 系统角色需求ER图": role_image,
            "图三-1 AI核心数据ER图": ai_image,
        },
    )

    print("567", result_567)
    print("v14", result_v14)


if __name__ == "__main__":
    main()
