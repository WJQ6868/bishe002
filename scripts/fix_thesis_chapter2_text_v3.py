from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm
from docx.text.paragraph import Paragraph


USECASE_IMAGE = Path(r"D:\bishe\one\generated\thesis_assets\fig_2_1_usecase_uml_black.png")


def get_input_doc() -> Path:
    desktop = Path.home() / "Desktop"
    matches = list(desktop.glob("2405273202*chapter2-updated-v2.docx"))
    if not matches:
        raise FileNotFoundError("Cannot find chapter2-updated-v2 document.")
    return matches[0]


def get_output_doc(input_doc: Path) -> Path:
    return input_doc.with_name(input_doc.stem.replace("-v2", "-v3") + input_doc.suffix)


def find_after(doc: Document, text: str, *, style_name: str | None = None, after_idx: int = -1) -> tuple[int, Paragraph]:
    for idx, para in enumerate(doc.paragraphs):
        if idx <= after_idx:
            continue
        if text in para.text and (style_name is None or para.style.name == style_name):
            return idx, para
    raise ValueError(f"Cannot find paragraph: {text}")


def delete_paragraph(paragraph: Paragraph) -> None:
    element = paragraph._element
    element.getparent().remove(element)


def insert_before(anchor: Paragraph, text: str = "", style=None) -> Paragraph:
    para = anchor.insert_paragraph_before(text)
    if style is not None:
        para.style = style
    return para


def replace_range_with_paragraphs(doc: Document, start_idx: int, end_idx: int, texts: list[str], style) -> None:
    anchor = doc.paragraphs[end_idx]
    targets = doc.paragraphs[start_idx:end_idx]
    for para in targets:
        delete_paragraph(para)
    for text in texts:
        insert_before(anchor, text, style)


def ensure_usecase_inserted(doc: Document, anchor: Paragraph, image_style, caption_style) -> None:
    for para in doc.paragraphs:
        if "系统角色需求用例图" in para.text:
            return
    image_para = insert_before(anchor, style=image_style)
    image_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    image_para.add_run().add_picture(str(USECASE_IMAGE), width=Cm(16))
    caption_para = insert_before(anchor, "图二-1 系统角色需求用例图", caption_style)
    caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER


def main() -> None:
    input_doc = get_input_doc()
    output_doc = get_output_doc(input_doc)
    doc = Document(input_doc)

    chapter2_idx, _ = find_after(doc, "需求分析与总体设计", style_name="Heading 1")
    sec21_idx, _ = find_after(doc, "业务场景与角色需求分析", style_name="Heading 2", after_idx=chapter2_idx)
    table21_idx, table21_para = find_after(doc, "角色需求与AI价值对应关系", style_name="表格标题", after_idx=sec21_idx)
    sec22_idx, _ = find_after(doc, "功能需求分析", style_name="Heading 2", after_idx=table21_idx)
    sec23_idx, _ = find_after(doc, "非功能需求分析", style_name="Heading 2", after_idx=sec22_idx)

    normal_style = doc.styles["Normal"]
    pic_caption_style = doc.paragraphs[table21_idx + 19].style if len(doc.paragraphs) > table21_idx + 19 else doc.styles["Normal"]
    # Use the style of the existing picture paragraph before architecture caption.
    arch_caption_idx, arch_caption_para = find_after(doc, "系统总体架构图", style_name="图片标题", after_idx=sec23_idx)
    image_style = doc.paragraphs[arch_caption_idx - 1].style
    caption_style = arch_caption_para.style

    texts_21 = [
        "第二章的需求分析不是先假定系统应该具备哪些AI能力，再去为这些能力寻找使用场景，而是从现有教务系统的真实使用过程反向推导出来的。学生、教师和管理员共用同一平台，但他们在系统中的任务目标、关注重点和停留页面并不一致。也正因为这种差异存在，系统在引入AI时不能简单采用“一套功能服务所有人”的思路，而必须先解释每类角色为什么会形成当前这些需求。",
        "学生侧需求最先暴露出来的问题，其实不是业务功能缺失，而是业务解释不足。传统教务系统通常已经提供了课程查询、请假申请、办事大厅、成绩查看等页面入口，学生真正感到困难的时刻，往往发生在看到入口之后。例如流程该从哪一步开始、不同事项需要准备什么材料、页面上的字段意味着什么、课程资料应该怎么理解，这些问题靠静态说明很难完全覆盖。因此，学生侧对AI的需求，本质上是希望系统能够在原有业务页面之上继续承担解释工作。",
        "从学生的操作链条看，这种需求并不是孤立产生的。学生进入系统后，常常先做查询，再做判断，最后才继续办理或学习。如果查询结果只能“看到”，却不能“看懂”，那么后续办理效率并不会真正提高。同样，在课程学习场景中，学生遇到的问题也不是单纯缺少资料，而是缺少能够结合课程资料进行即时解释的支持。这说明学生侧需要的不是独立聊天工具，而是嵌入在查询、办理和学习过程中的辅助能力。",
        "教师侧需求形成的原因又有所不同。教师在教务系统中的一类工作是确定性的，例如课程管理、成绩录入、审批处理和作业组织；另一类工作则更依赖教学内容本身，如整理课程资料、反复回答共性问题、准备教案和教学提纲。前一类工作传统教务系统已经能够较好支撑，后一类工作却长期依赖人工完成，重复频次高、零散时间多、复用效率低。这正是教师侧AI需求出现的现实基础。",
        "进一步分析教师的教学活动可以发现，教师并不会接受一个完全脱离课程边界的通用生成工具。课程答疑需要基于本课程资料，教案生成需要回到既定教学目标和教学内容，生成结果还必须允许教师继续调整。这意味着教师侧之所以会形成课程资料沉淀、课程助手问答和智能教案生成等需求，并不是单纯追求“自动化”，而是因为教学责任要求AI输出必须可追溯、可编辑、可复用。",
        "管理员侧需求则来自系统治理压力。只要AI能力真正进入教务平台，管理员很快就会面对一组新的问题：模型是否可用、知识资料由谁维护、不同场景是否应该共用同一套工作流、配置修改后前台是否能够及时同步。这些问题和学生、教师的直接使用体验不同，但它们决定了AI功能能否从演示能力转化为稳定能力。因此，管理员的需求并非附属需求，而是系统能否持续运行的前提条件。",
        "把三类角色放在同一条业务线上看，需求之间还存在明显的前后依赖。学生侧不断产生咨询与理解需求，教师侧不断补充课程资料和教学内容，管理员侧则负责保证模型、知识库和工作流处于可配置、可控制、可维护状态。没有教师资料，课程问答就缺少依据；没有管理员治理，前台的AI入口就难以长期稳定；没有学生与教师的真实使用，AI配置又无法沉淀为有效经验。需求分析正是从这种相互依赖关系中逐步抽象出来的。",
        "基于以上分析，系统最终可以归纳出三条较为稳定的需求主线：第一，学生侧需要在既有教务流程上获得即时解释和持续引导；第二，教师侧需要围绕课程资料形成可复用的教学支持能力；第三，管理员侧需要一套能够统一配置并稳定治理的AI资源组织方式。图二-1对这种“角色—需求—场景”的对应关系进行了更加直观的展示。",
    ]
    replace_range_with_paragraphs(doc, sec21_idx + 1, table21_idx, texts_21, normal_style)

    table21_idx, table21_para = find_after(doc, "角色需求与AI价值对应关系", style_name="表格标题", after_idx=sec21_idx)
    ensure_usecase_inserted(doc, table21_para, image_style, caption_style)

    sec22_idx, _ = find_after(doc, "功能需求分析", style_name="Heading 2", after_idx=table21_idx)
    sec23_idx, _ = find_after(doc, "非功能需求分析", style_name="Heading 2", after_idx=sec22_idx)

    texts_22 = [
        "功能需求并不是从技术模块出发拆分出来的，而是从前一节已经明确的角色压力和业务链条中一步步抽象出来的。换句话说，本系统为什么需要这些功能，不是因为AI能够做这些事，而是因为学生、教师和管理员在现有教务流程中已经暴露出对应的问题，系统需要用较小的结构改动去把这些问题接住。",
        "首先，基础教务功能必须保留为系统的主骨架。这一点并不是保守设计，而是需求本身决定的。学生的选课、请假、成绩查询、证书查看和办事事项办理，教师的课程管理、作业管理和成绩处理，管理员的用户维护和流程配置，构成了整个系统的真实业务边界。AI只有依附在这些边界之内，回答和生成的内容才有明确场景，不会脱离教务流程空转。",
        "其次，AI客服的功能需求来自高频但分散的咨询场景。学生和教师遇到的问题，很多并不值得专门找人处理，却又确实会影响后续操作，例如事项该找哪个入口、材料是否齐全、某项流程先后顺序是什么、规则解释如何理解。由此可以推导出，AI客服不能只是“能回答”，还需要具备推荐问题、状态反馈、知识不足提示和业务入口引导等功能，才能真正服务于咨询场景，而不是停留在泛化问答层面。",
        "AI课程助手的功能需求则是从教学答疑链中推出来的。教师在课程教学中经常面对重复问题，但这些问题又不能完全交给通用知识回答，因为不同课程、不同资料、不同教学重点之间差异很大。因此，系统需要支持教师围绕课程上传资料、形成课程专属知识来源，并让学生在明确的课程入口下发起提问。课程绑定、资料更新、助手选择和结果反馈这些功能，看起来是页面操作，实际上都是从“答疑必须贴合具体课程”这一需求前提中自然派生出来的。",
        "智能教案的功能需求来源于备课场景的重复性和迭代性。教师真正需要的不是一次性吐出一段文本，而是能围绕课程主题、资料基础和教学目标形成一份初稿，并在后续继续修改、保存和导出。正因为备课本身是一个多轮调整过程，系统才需要把教案生成设计成带有任务状态、结果回写、历史保留和再次编辑能力的功能，而不是把它处理成一次性的即时回答。",
        "管理端AI治理功能同样不是附属功能，而是由平台运行需求倒推出来的。模型接口管理、知识库管理、工作流编排、客服参数维护以及连通性校验，分别对应“能不能接入”“资料由谁维护”“不同场景如何分流”“最终展示什么内容”“配置改完能不能立即验证”这些很实际的问题。少了这一层，前台的AI入口很难长期稳定使用。",
        "除此之外，系统还需要一组跨模块的支撑功能。这些功能表面上不如AI客服、课程助手和智能教案显眼，但实际上决定了整个链条是否能闭合，包括角色权限校验、上传与生成状态提示、历史记录保存、结果再次复用以及异常信息反馈等。它们并不是额外附加的“优化项”，而是从完整业务链条中倒推出来的必需能力。也正因为有了这些支撑功能，系统中的AI模块才不是孤立页面，而是能与现有教务业务协同运行的整体能力。",
    ]
    replace_range_with_paragraphs(doc, sec22_idx + 1, sec23_idx, texts_22, normal_style)

    if output_doc.exists():
        output_doc.unlink()
    doc.save(output_doc)
    print(output_doc)


if __name__ == "__main__":
    main()
