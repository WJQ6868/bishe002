from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm


USECASE_IMAGE = Path(r"D:\bishe\one\generated\thesis_assets\fig_2_1_usecase_uml_black.png")


def main() -> None:
    desktop = Path.home() / "Desktop"
    input_doc = next(desktop.glob("2405273202*chapter2-updated-v3.docx"))
    output_doc = input_doc.with_name(input_doc.stem.replace("-v3", "-v4") + input_doc.suffix)

    doc = Document(input_doc)

    chapter2_idx = next(i for i, p in enumerate(doc.paragraphs) if p.style.name == "Heading 1" and p.text.strip() == "需求分析与总体设计")
    table_caption_idx = next(
        i
        for i, p in enumerate(doc.paragraphs)
        if i > chapter2_idx and p.style.name == "表格标题" and "角色需求与AI价值对应关系" in p.text
    )
    table_caption = doc.paragraphs[table_caption_idx]

    for p in doc.paragraphs:
        if "系统角色需求用例图" in p.text:
            if output_doc.exists():
                output_doc.unlink()
            doc.save(output_doc)
            print(output_doc)
            return

    arch_caption_idx = next(
        i
        for i, p in enumerate(doc.paragraphs)
        if i > chapter2_idx and p.style.name == "图片标题" and "系统总体架构图" in p.text
    )
    image_style = doc.paragraphs[arch_caption_idx - 1].style
    caption_style = doc.paragraphs[arch_caption_idx].style

    img_para = table_caption.insert_paragraph_before("")
    img_para.style = image_style
    img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    img_para.add_run().add_picture(str(USECASE_IMAGE), width=Cm(16))

    cap_para = table_caption.insert_paragraph_before("图二-1 系统角色需求用例图")
    cap_para.style = caption_style
    cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if output_doc.exists():
        output_doc.unlink()
    doc.save(output_doc)
    print(output_doc)


if __name__ == "__main__":
    main()
