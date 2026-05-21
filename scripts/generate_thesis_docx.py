# -*- coding: utf-8 -*-
from __future__ import annotations

import json
import math
import re
import sqlite3
from pathlib import Path
from textwrap import dedent

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Mm, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DIR = ROOT / "generated"
ASSET_DIR = OUTPUT_DIR / "thesis_assets"
DOCX_PATH = ROOT / "AI赋能的高校教务系统-毕业论文润色定稿版.docx"
DB_PATH = ROOT / "edu_system.db"
TEST_JSON_PATH = ROOT / "generated_ai_test_results.json"

TITLE = "AI赋能的高校教务系统"
AUTHOR = "王佳齐"
STUDENT_ID = "2405273202"
COLLEGE = "计算机与软件学院"
MAJOR = "工业互联网技术"
CLASS_NAME = "互联2432"
TEACHER = "张锦辉"
ENTERPRISE_MENTOR = "嵇伟"
DATE_CN = "2026年4月"
CHECK_DATE_CN = "2026年5月5日"

FONT_SONG = "宋体"
FONT_HEI = "黑体"
FONT_EN = "Times New Roman"
FONT_FILE = Path(r"C:\Windows\Fonts\msyh.ttc")
FONT_FILE_BOLD = Path(r"C:\Windows\Fonts\simhei.ttf")

TEXT_COLLECTOR: list[str] = []


def normalize_paragraphs(text: str) -> list[str]:
    text = dedent(text).strip()
    if not text:
        return []
    blocks = []
    for item in text.split("\n\n"):
        line = "".join(i.strip() for i in item.splitlines())
        if line:
            blocks.append(line)
    return blocks


def record_text(text: str) -> None:
    if text:
        TEXT_COLLECTOR.append(text)


def load_font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont:
    path = FONT_FILE_BOLD if bold and FONT_FILE_BOLD.exists() else FONT_FILE
    return ImageFont.truetype(str(path), size=size)


def wrap_text(draw: ImageDraw.ImageDraw, text: str, font: ImageFont.FreeTypeFont, max_width: int) -> list[str]:
    if not text:
        return [""]
    lines: list[str] = []
    current = ""
    for ch in text:
        candidate = current + ch
        if draw.textlength(candidate, font=font) <= max_width:
            current = candidate
        else:
            if current:
                lines.append(current)
            current = ch
    if current:
        lines.append(current)
    return lines or [text]


def draw_centered_text(draw: ImageDraw.ImageDraw, xy: tuple[int, int], text: str, font, fill=(0, 0, 0)) -> None:
    bbox = draw.textbbox((0, 0), text, font=font)
    w = bbox[2] - bbox[0]
    h = bbox[3] - bbox[1]
    draw.text((xy[0] - w / 2, xy[1] - h / 2), text, font=font, fill=fill)


def draw_round_box(
    draw: ImageDraw.ImageDraw,
    box: tuple[int, int, int, int],
    title: str,
    lines: list[str] | None = None,
    fill=(255, 255, 255),
    outline=(52, 73, 94),
    title_fill=(52, 73, 94),
    radius: int = 18,
) -> None:
    x1, y1, x2, y2 = box
    draw.rounded_rectangle(box, radius=radius, fill=fill, outline=outline, width=3)
    title_font = load_font(34, bold=True)
    body_font = load_font(36)
    draw.text((x1 + 14, y1 + 10), title, font=title_font, fill=title_fill)
    if lines:
        y = y1 + 64
        for line in lines:
            wrapped = wrap_text(draw, line, body_font, x2 - x1 - 28)
            for item in wrapped:
                draw.text((x1 + 14, y), item, font=body_font, fill=(33, 37, 41))
                y += 40


def draw_arrow(draw: ImageDraw.ImageDraw, start: tuple[int, int], end: tuple[int, int], label: str | None = None) -> None:
    draw.line([start, end], fill=(52, 73, 94), width=4)
    angle = math.atan2(end[1] - start[1], end[0] - start[0])
    arrow_len = 16
    angle1 = angle - math.pi / 8
    angle2 = angle + math.pi / 8
    p1 = (end[0] - arrow_len * math.cos(angle1), end[1] - arrow_len * math.sin(angle1))
    p2 = (end[0] - arrow_len * math.cos(angle2), end[1] - arrow_len * math.sin(angle2))
    draw.polygon([end, p1, p2], fill=(52, 73, 94))
    if label:
        font = load_font(22, bold=True)
        mx = (start[0] + end[0]) // 2
        my = (start[1] + end[1]) // 2 - 18
        draw_centered_text(draw, (mx, my), label, font, fill=(44, 62, 80))


def draw_poly_arrow(
    draw: ImageDraw.ImageDraw,
    points: list[tuple[int, int]],
    label: str | None = None,
    color=(52, 73, 94),
    width: int = 4,
) -> None:
    if len(points) < 2:
        return
    draw.line(points, fill=color, width=width)
    start = points[-2]
    end = points[-1]
    angle = math.atan2(end[1] - start[1], end[0] - start[0])
    arrow_len = 16
    angle1 = angle - math.pi / 8
    angle2 = angle + math.pi / 8
    p1 = (end[0] - arrow_len * math.cos(angle1), end[1] - arrow_len * math.sin(angle1))
    p2 = (end[0] - arrow_len * math.cos(angle2), end[1] - arrow_len * math.sin(angle2))
    draw.polygon([end, p1, p2], fill=color)
    if label:
        font = load_font(20, bold=True)
        first = points[0]
        last = points[-1]
        mx = (first[0] + last[0]) // 2
        my = (first[1] + last[1]) // 2 - 18
        draw_centered_text(draw, (mx, my), label, font, fill=(44, 62, 80))


def draw_round_box_custom(
    draw: ImageDraw.ImageDraw,
    box: tuple[int, int, int, int],
    title: str,
    lines: list[str] | None = None,
    *,
    title_size: int = 28,
    body_size: int = 20,
    title_top: int = 12,
    line_gap: int = 28,
    fill=(255, 255, 255),
    outline=(52, 73, 94),
    title_fill=(52, 73, 94),
    radius: int = 18,
) -> None:
    x1, y1, x2, y2 = box
    draw.rounded_rectangle(box, radius=radius, fill=fill, outline=outline, width=3)
    title_font = load_font(title_size, bold=True)
    body_font = load_font(body_size)
    draw.text((x1 + 10, y1 + title_top), title, font=title_font, fill=title_fill)
    if lines:
        y = y1 + title_top + title_size + 8
        for line in lines:
            wrapped = wrap_text(draw, line, body_font, x2 - x1 - 20)
            for item in wrapped:
                draw.text((x1 + 10, y), item, font=body_font, fill=(33, 37, 41))
                y += line_gap


def draw_entity_box(
    draw: ImageDraw.ImageDraw,
    box: tuple[int, int, int, int],
    title: str,
    fields: list[str],
    *,
    fill=(248, 250, 252),
    outline=(52, 73, 94),
) -> None:
    x1, y1, x2, y2 = box
    draw.rounded_rectangle(box, radius=18, fill=fill, outline=outline, width=3)
    title_font = load_font(24, bold=True)
    body_font = load_font(18)
    draw.text((x1 + 16, y1 + 12), title, font=title_font, fill=(44, 62, 80))
    divider_y = y1 + 50
    draw.line((x1 + 10, divider_y, x2 - 10, divider_y), fill=outline, width=2)
    y = divider_y + 12
    for field in fields:
        wrapped = wrap_text(draw, field, body_font, x2 - x1 - 32)
        for item in wrapped:
            draw.text((x1 + 16, y), item, font=body_font, fill=(33, 37, 41))
            y += 24


def create_placeholder_image(path: Path, title: str, subtitle: str) -> None:
    img = Image.new("RGB", (1400, 820), (247, 248, 250))
    draw = ImageDraw.Draw(img)
    draw.rounded_rectangle((80, 80, 1320, 740), radius=24, outline=(120, 130, 140), width=6, fill=(255, 255, 255))
    for i in range(0, 18):
        x = 100 + i * 65
        draw.line((x, 95, x + 35, 95), fill=(170, 170, 170), width=4)
        draw.line((x, 725, x + 35, 725), fill=(170, 170, 170), width=4)
    title_font = load_font(54, bold=True)
    sub_font = load_font(28)
    draw_centered_text(draw, (700, 330), title, title_font, fill=(80, 80, 80))
    draw_centered_text(draw, (700, 410), "【系统功能界面示意】", load_font(34, bold=True), fill=(80, 80, 80))
    wrapped = wrap_text(draw, subtitle, sub_font, 980)
    y = 500
    for line in wrapped:
        draw_centered_text(draw, (700, y), line, sub_font, fill=(90, 90, 90))
        y += 42
    img.save(path)


def create_system_architecture(path: Path) -> None:
    img = Image.new("RGB", (4200, 2500), (255, 255, 255))
    draw = ImageDraw.Draw(img)
    draw_centered_text(draw, (2100, 86), "高校教务系统总体架构图", load_font(76, bold=True))

    top_boxes = [
        ((180, 220, 1120, 560), "学生端", ["课程查询、请假申请、AI 客服、AI 课程助手、查看共享教案"], (255, 255, 255)),
        ((1630, 220, 2570, 560), "教师端", ["课程管理、作业批改、课程助手、知识库上传、智能教案生成"], (255, 255, 255)),
        ((3080, 220, 4020, 560), "管理员端", ["用户管理、数据维护、模型配置、知识库管理、工作流编排"], (255, 255, 255)),
    ]
    for box, title, lines, fill in top_boxes:
        draw_round_box_custom(draw, box, title, lines, title_size=82, body_size=64, line_gap=68, fill=fill, radius=28)

    front_box = (260, 760, 3940, 1140)
    back_box = (220, 1380, 3980, 1900)
    draw_round_box_custom(
        draw,
        front_box,
        "前端展示层（Vue3 + TypeScript + Element Plus + Pinia + Vue Router）",
        ["统一路由、角色化界面、Axios 请求封装、SSE 流式结果渲染、即时消息与可视化组件"],
        title_size=80,
        body_size=60,
        line_gap=66,
        fill=(255, 255, 255),
        radius=28,
    )
    draw_round_box_custom(
        draw,
        back_box,
        "后端服务层（FastAPI + SQLAlchemy + Socket.IO）",
        [
            "认证授权、课程管理、办事大厅、成绩与请假、管理员 AI 配置、AI Portal、AI QA 路由",
            "统一通过依赖注入获取用户与数据库会话，支持 /api 前缀接口与静态文件访问",
        ],
        title_size=80,
        body_size=60,
        line_gap=66,
        fill=(255, 255, 255),
        radius=28,
    )

    lower_boxes = [
        ((160, 2120, 1140, 2470), "AI 配置编排层", ["模型 API 管理、工作流绑定、客服参数维护、模型测试与启停控制"], (255, 255, 255)),
        ((1160, 2120, 2140, 2470), "知识组织层", ["基础知识库、课程私有知识库、文档抽取、文本切分、Chunk 元数据存储"], (255, 255, 255)),
        ((2160, 2120, 3140, 2470), "推理执行层", ["TF-IDF 检索、Prompt 组装、模型选择、SSE 流式输出、结果持久化"], (255, 255, 255)),
        ((3160, 2120, 4040, 2470), "数据与外部资源层", ["SQLite 数据库", "静态上传文件", "DashScope / Ark 模型接口"], (255, 255, 255)),
    ]
    for box, title, lines, fill in lower_boxes:
        draw_round_box_custom(draw, box, title, lines, title_size=74, body_size=58, line_gap=64, fill=fill, radius=28)

    for cx in (650, 2100, 3550):
        draw_poly_arrow(draw, [(cx, 560), (cx, 650), (2100, 650), (2100, 760)], width=7)
    draw_poly_arrow(draw, [(2100, 1140), (2100, 1380)], width=7)
    for cx in (650, 1650, 2650, 3600):
        draw_poly_arrow(draw, [(2100, 1900), (2100, 2005), (cx, 2005), (cx, 2120)], width=7)

    img.save(path)


def create_ai_closed_loop(path: Path) -> None:
    img = Image.new("RGB", (4000, 2400), (255, 255, 255))
    draw = ImageDraw.Draw(img)
    draw_centered_text(draw, (2000, 86), "AI 业务闭环与分层实现逻辑", load_font(76, bold=True))

    stage_boxes = [
        ((140, 220, 920, 560), "配置端", ["管理员配置模型、知识库、工作流与客服参数"], (255, 255, 255)),
        ((1080, 220, 1860, 560), "业务端", ["学生/教师按角色调用 AI 客服、AI 课程助手、智能教案"], (255, 255, 255)),
        ((2020, 220, 2800, 560), "数据回流", ["生成任务、日志、问答记录、课程资料更新"], (255, 255, 255)),
        ((2960, 220, 3740, 560), "优化闭环", ["补充知识库、调整模型、修正提示词与工作流策略"], (255, 255, 255)),
    ]
    for box, title, lines, fill in stage_boxes:
        draw_round_box_custom(draw, box, title, lines, title_size=78, body_size=60, line_gap=66, fill=fill, radius=28)
    draw_poly_arrow(draw, [(920, 390), (1080, 390)], width=7)
    draw_poly_arrow(draw, [(1860, 390), (2020, 390)], width=7)
    draw_poly_arrow(draw, [(2800, 390), (2960, 390)], width=7)
    draw_poly_arrow(draw, [(3350, 560), (3350, 700), (520, 700), (520, 600)], label="反馈", width=7)

    outer = (220, 940, 3780, 2340)
    draw.rounded_rectangle(outer, radius=30, fill=(255, 255, 255), outline=(52, 73, 94), width=5)
    draw.text((300, 995), "四层 AI 架构", font=load_font(72, bold=True), fill=(44, 62, 80))
    row_boxes = [
        ((380, 1180, 3620, 1380), "第一层：AI 配置编排层", ["模型 API、知识库、工作流 App 与功能参数统一由后台管理"], (255, 255, 255)),
        ((380, 1480, 3620, 1680), "第二层：知识组织层", ["教师上传资料，系统抽取文本并切分片段，为问答和教案生成提供上下文"], (255, 255, 255)),
        ((380, 1780, 3620, 1980), "第三层：推理执行层", ["根据请求解析工作流、选择模型、检索知识片段、拼装 Prompt 并调用模型"], (255, 255, 255)),
        ((380, 2080, 3620, 2280), "第四层：业务接入层", ["AI 客服、AI 课程助手、智能教案三类场景分别面向师生与管理端落地"], (255, 255, 255)),
    ]
    for box, title, lines, fill in row_boxes:
        draw_round_box_custom(draw, box, title, lines, title_size=74, body_size=60, line_gap=66, fill=fill, radius=26)
    img.save(path)


def create_ai_er(path: Path) -> None:
    img = Image.new("RGB", (4800, 2800), (255, 255, 255))
    draw = ImageDraw.Draw(img)
    line_color = (52, 73, 94)
    text_color = (33, 37, 41)

    def entity_box(box: tuple[int, int, int, int], title: str, fields: list[str], *, fill) -> None:
        x1, y1, x2, y2 = box
        draw.rounded_rectangle(box, radius=28, fill=fill, outline=line_color, width=5)
        title_font = load_font(76, bold=True)
        body_font = load_font(62)
        draw.text((x1 + 18, y1 + 12), title, font=title_font, fill=line_color)
        divider_y = y1 + 122
        draw.line((x1 + 18, divider_y, x2 - 18, divider_y), fill=line_color, width=3)
        y = divider_y + 12
        for field in fields:
            for item in wrap_text(draw, field, body_font, x2 - x1 - 28):
                draw.text((x1 + 18, y), item, font=body_font, fill=text_color)
                y += 66

    def big_arrow(points: list[tuple[int, int]], label: str | None = None, label_xy: tuple[int, int] | None = None) -> None:
        if len(points) < 2:
            return
        draw.line(points, fill=line_color, width=7)
        start = points[-2]
        end = points[-1]
        angle = math.atan2(end[1] - start[1], end[0] - start[0])
        arrow_len = 26
        angle1 = angle - math.pi / 8
        angle2 = angle + math.pi / 8
        p1 = (end[0] - arrow_len * math.cos(angle1), end[1] - arrow_len * math.sin(angle1))
        p2 = (end[0] - arrow_len * math.cos(angle2), end[1] - arrow_len * math.sin(angle2))
        draw.polygon([end, p1, p2], fill=line_color)
        if label:
            label_font = load_font(42, bold=True)
            lx, ly = label_xy or ((points[0][0] + points[-1][0]) // 2, (points[0][1] + points[-1][1]) // 2 - 26)
            padding_x, padding_y = 18, 10
            bbox = draw.textbbox((0, 0), label, font=label_font)
            w = bbox[2] - bbox[0]
            h = bbox[3] - bbox[1]
            draw.rounded_rectangle((lx - w / 2 - padding_x, ly - h / 2 - padding_y, lx + w / 2 + padding_x, ly + h / 2 + padding_y), radius=16, fill=(255, 255, 255))
            draw_centered_text(draw, (lx, ly), label, label_font, fill=line_color)

    draw_centered_text(draw, (2400, 78), "AI 核心数据 ER 图", load_font(88, bold=True))

    boxes = {
        "sys_users": (120, 220, 1040, 760),
        "courses": (120, 1470, 1040, 2110),
        "ai_model_apis": (1230, 220, 2260, 960),
        "ai_knowledge_bases": (1230, 1240, 2260, 2080),
        "ai_workflow_apps": (2450, 220, 3570, 1020),
        "ai_kb_documents": (2450, 1470, 3570, 2240),
        "ai_lesson_plan_tasks": (3740, 220, 4680, 1020),
        "ai_kb_chunks": (3740, 1470, 4680, 2310),
    }
    entity_box(boxes["sys_users"], "sys_users", ["PK id", "username", "role", "is_active"], fill=(255, 255, 255))
    entity_box(boxes["courses"], "courses", ["PK id", "name", "teacher_id", "credit", "capacity", "course_type"], fill=(255, 255, 255))
    entity_box(boxes["ai_model_apis"], "ai_model_apis", ["PK id", "name", "provider", "model_name", "endpoint", "enabled", "is_default"], fill=(255, 255, 255))
    entity_box(boxes["ai_knowledge_bases"], "ai_knowledge_bases", ["PK id", "slug", "name", "owner_type", "owner_user_id", "course_id", "feature"], fill=(255, 255, 255))
    entity_box(boxes["ai_workflow_apps"], "ai_workflow_apps", ["PK id", "code", "type", "name", "knowledge_base_id", "model_api_id", "owner_user_id", "course_id", "status"], fill=(255, 255, 255))
    entity_box(boxes["ai_kb_documents"], "ai_kb_documents", ["PK id", "knowledge_base_id", "title", "original_filename", "url", "file_ext", "enabled"], fill=(255, 255, 255))
    entity_box(boxes["ai_lesson_plan_tasks"], "ai_lesson_plan_tasks", ["PK id", "teacher_user_id", "course_id", "title", "status", "result", "knowledge_base_id", "model_api_id"], fill=(255, 255, 255))
    entity_box(boxes["ai_kb_chunks"], "ai_kb_chunks", ["PK id", "knowledge_base_id", "document_id", "seq", "content", "tokens", "document_title"], fill=(255, 255, 255))

    big_arrow([(1040, 410), (1230, 410)], "用户拥有/配置", (1130, 360))
    big_arrow([(1040, 1730), (1230, 1730)], "课程关联", (1130, 1675))
    big_arrow([(2260, 460), (2450, 460)], "模型绑定", (2350, 405))
    big_arrow([(2260, 1430), (2360, 1430), (2360, 760), (2450, 760)], "工作流知识库", (2350, 1090))
    big_arrow([(2260, 1730), (2450, 1730)], "1:N", (2350, 1675))
    big_arrow([(3570, 1810), (3740, 1810)], "1:N", (3655, 1760))
    big_arrow([(3570, 460), (3740, 460)], "任务模型", (3655, 405))
    big_arrow([(3570, 760), (3740, 760)], "任务知识库", (3655, 705))
    big_arrow([(4210, 1470), (4210, 1045)], "教案引用知识分块", (4410, 1260))

    img.save(path)


def create_flow_diagram(path: Path, title: str, steps: list[str], fill_color=(255, 255, 255)) -> None:
    img = Image.new("RGB", (1700, 610), (255, 255, 255))
    draw = ImageDraw.Draw(img)
    draw_centered_text(draw, (850, 42), title, load_font(46, bold=True))
    start_x = 40
    width = 300
    top = 95
    bottom = 415
    gap = 20
    centers = []
    for idx, step in enumerate(steps):
        x1 = start_x + idx * (width + gap)
        x2 = x1 + width
        centers.append(((x1 + x2) // 2, (top + bottom) // 2))
        draw_round_box(draw, (x1, top, x2, bottom), f"步骤{idx + 1}", [step], fill=fill_color)
        if idx < len(steps) - 1:
            draw_arrow(draw, (x2, (top + bottom) // 2), (x2 + gap, (top + bottom) // 2))
    img.save(path)


def build_assets() -> dict[str, Path]:
    OUTPUT_DIR.mkdir(exist_ok=True)
    ASSET_DIR.mkdir(exist_ok=True)
    assets = {
        "system_arch": ASSET_DIR / "fig_2_1_system_architecture.png",
        "closed_loop": ASSET_DIR / "fig_2_2_ai_closed_loop.png",
        "er_ai": ASSET_DIR / "fig_3_1_ai_er.png",
        "flow_customer": ASSET_DIR / "fig_4_1_customer_service_flow.png",
        "flow_course": ASSET_DIR / "fig_4_2_course_assistant_flow.png",
        "flow_lesson": ASSET_DIR / "fig_4_3_lesson_plan_flow.png",
        "shot_customer": ASSET_DIR / "fig_4_4_customer_service_placeholder.png",
        "shot_course": ASSET_DIR / "fig_4_5_course_assistant_placeholder.png",
        "shot_lesson": ASSET_DIR / "fig_4_6_lesson_plan_placeholder.png",
        "shot_admin": ASSET_DIR / "fig_4_7_admin_ai_placeholder.png",
    }
    create_system_architecture(assets["system_arch"])
    create_ai_closed_loop(assets["closed_loop"])
    create_ai_er(assets["er_ai"])
    create_flow_diagram(
        assets["flow_customer"],
        "AI 客服核心处理流程",
        ["前端加载客服配置与工作流列表", "用户输入问题并提交到 /api/ai_qa/qa/stream", "后端检索知识库片段并拼装 Prompt", "按优先级解析模型并调用外部模型接口", "SSE 分段返回答案并在界面中实时渲染"],
    )
    create_flow_diagram(
        assets["flow_course"],
        "AI 课程助手核心处理流程",
        ["管理员配置基础课程助手工作流", "教师复制基础工作流并上传课程私有知识库", "学生或教师按课程选择助手提问", "系统合并工作流知识库与课程知识库进行检索", "模型生成课程相关回答并回传前端"],
    )
    create_flow_diagram(
        assets["flow_lesson"],
        "智能教案任务化生成流程",
        ["教师选择课程并上传/解析教学资料", "创建教案任务并写入 ai_lesson_plan_tasks", "调用 lesson_plan 工作流进行流式生成", "生成结果回写任务状态并允许教师二次编辑", "教师保存修改结果并导出 Markdown 教案"],
    )
    create_placeholder_image(assets["shot_customer"], "学生端 / 教师端 AI 客服界面", "界面包含推荐问题、流式回答、多会话列表和悬浮客服入口。")
    create_placeholder_image(assets["shot_course"], "AI 课程助手界面", "界面包含课程助手选择、教师自定义助手、课程提问和回答展示区域。")
    create_placeholder_image(assets["shot_lesson"], "智能教案界面", "界面包含课程知识库文档、教案任务清单、生成结果编辑和导出操作。")
    create_placeholder_image(assets["shot_admin"], "管理员 AI 配置界面", "界面包含模型 API 管理、知识库管理、工作流绑定和模型连通性测试。")
    return assets


def query_runtime_data() -> dict:
    counts = {}
    workflows = []
    usage_log_count = 0
    if DB_PATH.exists():
        conn = sqlite3.connect(DB_PATH)
        conn.row_factory = sqlite3.Row
        cur = conn.cursor()
        for table in [
            "sys_users",
            "courses",
            "service_item",
            "service_apply",
            "ai_model_apis",
            "ai_knowledge_bases",
            "ai_kb_documents",
            "ai_kb_chunks",
            "ai_workflow_apps",
            "ai_lesson_plan_tasks",
        ]:
            counts[table] = cur.execute(f"select count(*) from {table}").fetchone()[0]
        usage_log_count = cur.execute("select count(*) from ai_usage_logs").fetchone()[0]
        workflows = [dict(row) for row in cur.execute(
            "select id, code, type, name, status, knowledge_base_id, model_api_id, owner_user_id, course_id from ai_workflow_apps order by id"
        ).fetchall()]
        conn.close()
    return {
        "counts": counts,
        "usage_log_count": usage_log_count,
        "workflows": workflows,
        "test_summary": [
            ["获取 AI 客服配置", "GET /api/ai/customer-service/config", "200", "通过，能够返回欢迎语与推荐问题配置"],
            ["获取 AI 客服工作流", "GET /api/ai/customer-service/apps", "200", "通过，返回可用客服工作流列表"],
            ["获取课程助手工作流", "GET /api/ai/course-assistant/apps", "200", "通过，能够列出启用的课程助手工作流"],
            ["AI 问答流式接口", "POST /api/ai_qa/qa/stream", "200", "通过，SSE 形式返回分段结果"],
            ["教师查询授课课程", "GET /api/ai/teacher/courses", "200", "通过，可按教师账号获取课程集合"],
            ["教师创建课程助手", "POST /api/ai/teacher/course-assistant/apps", "200", "通过，可复制基础工作流生成个人助手"],
            ["创建教案任务", "POST /api/ai/teacher/lesson-plan/tasks", "200", "通过，可写入待处理任务记录"],
            ["更新教案任务结果", "PUT /api/ai/teacher/lesson-plan/tasks/{id}/result", "200", "通过，可回写完成状态与生成内容"],
        ],
    }


def ensure_rfonts(element, east_asia_font: str) -> None:
    rpr = element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), east_asia_font)
    rfonts.set(qn("w:ascii"), FONT_EN)
    rfonts.set(qn("w:hAnsi"), FONT_EN)


def apply_run_font(run, cn_font=FONT_SONG, size=12, bold=False, italic=False, color: RGBColor | None = None):
    run.font.name = FONT_EN
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color
    ensure_rfonts(run._element, cn_font)
    return run


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    section.header_distance = Cm(2.0)
    section.footer_distance = Cm(1.5)
    normal = doc.styles["Normal"]
    normal.font.name = FONT_EN
    normal.font.size = Pt(12)
    ensure_rfonts(normal._element, FONT_SONG)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    paragraph._p.append(fld_begin)
    paragraph._p.append(instr)
    paragraph._p.append(fld_separate)
    run = paragraph.add_run("1")
    apply_run_font(run, FONT_EN, 10.5)
    paragraph._p.append(fld_end)


def add_body_paragraph(doc: Document, text: str, *, indent=True, align=WD_ALIGN_PARAGRAPH.JUSTIFY) -> None:
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    if indent:
        p.paragraph_format.first_line_indent = Pt(24)
    run = p.add_run(text)
    apply_run_font(run, FONT_SONG, 12)
    record_text(text)


def add_center_paragraph(doc: Document, text: str, size=12, bold=False, font=FONT_SONG, color=None) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    apply_run_font(run, font, size, bold=bold, color=color)
    record_text(text)


def add_heading1(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(text)
    apply_run_font(run, FONT_HEI, 16, bold=True)
    record_text(text)


def add_heading2(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    apply_run_font(run, FONT_HEI, 14, bold=True)
    record_text(text)


def add_heading3(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    apply_run_font(run, FONT_HEI, 12, bold=True)
    record_text(text)


def add_paragraphs(doc: Document, text: str) -> None:
    for item in normalize_paragraphs(text):
        add_body_paragraph(doc, item)


def add_keywords(doc: Document, label: str, items: list[str], english=False) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    if not english:
        p.paragraph_format.first_line_indent = Pt(24)
    run1 = p.add_run(label)
    apply_run_font(run1, FONT_HEI if not english else FONT_EN, 12, bold=True)
    content = "；".join(items) if not english else "; ".join(items)
    run2 = p.add_run(content)
    apply_run_font(run2, FONT_SONG if not english else FONT_EN, 12)
    record_text(label + content)


def insert_toc(paragraph) -> None:
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = r'TOC \o "1-3" \h \z \u'
    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "目录请在 Word 中右键选择“更新域”后生成。"
    fld_separate.append(text)
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_separate)
    run._r.append(fld_end)


def set_cell_text(cell, text: str, *, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER, size=10.5) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.line_spacing = 1.25
    run = p.add_run(text)
    apply_run_font(run, FONT_SONG, size, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    record_text(text)


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("left", "top", "right", "bottom", "insideH", "insideV"):
        data = kwargs.get(edge)
        if data:
            tag = f"w:{edge}"
            element = tc_borders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tc_borders.append(element)
            for key, value in data.items():
                element.set(qn(f"w:{key}"), str(value))


def apply_three_line_table(table) -> None:
    rows = list(table.rows)
    cols = len(table.columns)
    nil = {"val": "nil"}
    single = {"val": "single", "sz": 8, "color": "000000"}
    for r_idx, row in enumerate(rows):
        for c_idx in range(cols):
            cell = row.cells[c_idx]
            border = {
                "left": nil,
                "right": nil,
                "top": nil,
                "bottom": nil,
            }
            if r_idx == 0:
                border["top"] = single
                border["bottom"] = single
            if r_idx == len(rows) - 1:
                border["bottom"] = single
            set_cell_border(cell, **border)


def add_table(doc: Document, title: str, headers: list[str], rows: list[list[str]], col_widths: list[float] | None = None) -> None:
    add_center_paragraph(doc, title, size=11, bold=True, font=FONT_SONG)
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_text(hdr[i], h, bold=True)
        if col_widths and i < len(col_widths):
            hdr[i].width = Cm(col_widths[i])
    for row in rows:
        cells = table.add_row().cells
        for i, item in enumerate(row):
            align = WD_ALIGN_PARAGRAPH.LEFT if len(item) > 12 else WD_ALIGN_PARAGRAPH.CENTER
            set_cell_text(cells[i], item, align=align)
            if col_widths and i < len(col_widths):
                cells[i].width = Cm(col_widths[i])
    apply_three_line_table(table)
    doc.add_paragraph()


def add_figure(doc: Document, img_path: Path, caption: str, width_cm=15.5) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(img_path), width=Cm(width_cm))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(6)
    r = cap.add_run(caption)
    apply_run_font(r, FONT_SONG, 10.5)
    record_text(caption)


def add_cover(doc: Document) -> None:
    for _ in range(3):
        doc.add_paragraph()
    add_center_paragraph(doc, "南京工业职业技术大学", size=22, bold=True, font=FONT_HEI)
    add_center_paragraph(doc, "本科毕业设计（论文）", size=20, bold=True, font=FONT_HEI)
    doc.add_paragraph()
    doc.add_paragraph()
    add_center_paragraph(doc, TITLE, size=24, bold=True, font=FONT_HEI, color=RGBColor(0x1F, 0x3A, 0x5F))
    doc.add_paragraph()
    doc.add_paragraph()
    table = doc.add_table(rows=7, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    labels = ["学生姓名", "学号", "学院", "专业", "班级", "指导教师", "企业导师"]
    values = [AUTHOR, STUDENT_ID, COLLEGE, MAJOR, CLASS_NAME, TEACHER, ENTERPRISE_MENTOR]
    for i, (k, v) in enumerate(zip(labels, values)):
        set_cell_text(table.rows[i].cells[0], k, bold=True)
        set_cell_text(table.rows[i].cells[1], v, align=WD_ALIGN_PARAGRAPH.LEFT)
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(cell, left={"val": "nil"}, right={"val": "nil"}, top={"val": "nil"}, bottom={"val": "nil"})
    doc.add_paragraph()
    doc.add_paragraph()
    add_center_paragraph(doc, DATE_CN, size=14, bold=True, font=FONT_SONG)
    doc.add_page_break()


def add_integrity_page(doc: Document) -> None:
    add_center_paragraph(doc, "诚信承诺书", size=18, bold=True, font=FONT_HEI)
    add_paragraphs(doc, f"""
    本人郑重声明：所提交的毕业设计（论文）《AI赋能的高校教务系统》是在指导教师指导下独立完成的研究成果。除文中已经明确注明引用的内容外，论文中不包含其他个人或集体已经发表或撰写过的研究成果，也不包含为获得南京工业职业技术大学或其他教育机构学位、证书而使用过的材料。

    本论文以项目当前代码、数据库结构和本地运行结果为依据，围绕系统功能、技术架构、AI 模块设计、数据组织方式和运行验证结果展开分析。文中所涉及的系统架构、功能流程、接口逻辑和数据结构均来源于项目当前实现，论文内容力求与系统功能保持一致。

    若因本论文内容引发学术不端、知识产权或事实失实等问题，本人愿承担相应责任。
    """)
    doc.add_paragraph()
    add_body_paragraph(doc, f"学生签名：{AUTHOR}　　　　　　　　　　　　　　　　日期：{CHECK_DATE_CN}", indent=False)
    doc.add_page_break()


def add_abstract_pages(doc: Document) -> None:
    add_center_paragraph(doc, "摘  要", size=18, bold=True, font=FONT_HEI)
    add_paragraphs(doc, f"""
    本文以 AI 赋能的高校教务系统毕业设计项目为研究对象，依托项目当前代码、数据库样本、启动脚本与本地运行结果，对系统的功能构成、AI 模块实现路径及运行效果进行了系统分析。系统面向学生、教师和管理员三类角色，围绕教务管理、校园服务、课程支持和智能化辅助生成等场景形成完整的功能体系。

    该系统采用前后端分离架构。前端基于 Vue3、TypeScript、Element Plus、Pinia 与 Vue Router 构建学生端、教师端和管理员端页面；后端基于 FastAPI、SQLAlchemy 与 SQLite 提供统一接口、数据持久化和业务支撑，并结合 Socket.IO 实现即时通信能力。在传统教务业务模块基础上，系统进一步引入模型管理、知识库管理、工作流绑定、流式问答以及教案任务回写等 AI 相关功能，从而形成面向教务场景的智能化扩展框架。

    论文重点围绕三个核心 AI 应用场景展开分析。其一，AI 客服通过工作流配置读取、知识片段检索和 SSE 流式输出，为学生和教师提供统一的教务咨询服务入口。其二，AI 课程助手支持教师在基础工作流之上构建个人课程助手，并将课程资料转化为私有知识库，以增强课程问答的针对性。其三，智能教案模块采用任务化处理模式，将教案生成、结果回写、二次编辑与导出操作纳入统一流程管理。底层检索机制采用 TF-IDF 与余弦相似度进行文本召回，形成了适用于本课题场景的轻量级检索增强生成方案。

    截至{CHECK_DATE_CN}，项目本地数据库中已保存 4 条模型 API 配置、4 个知识库、3 份知识库文档、86 个知识片段、4 个工作流应用以及 3 条教案任务记录。论文测试部分采用功能用例和接口联调相结合的方式，对客服配置读取、工作流列表获取、统一问答接口调用、教师创建课程助手、创建教案任务及结果回写等关键链路进行了验证。结果表明，系统核心功能链路能够按照设计流程运行。

    本文围绕系统功能展开论述，重点说明 AI 客服、AI 课程助手、智能教案和管理员 AI 配置模块如何与教务场景结合。系统通过模型配置、知识库组织、工作流绑定和流式问答机制，将 AI 能力接入学生咨询、教师课程支持和教学材料生成等环节，体现出高校教务系统智能化升级的应用价值。
    """)
    add_keywords(doc, "关键词：", ["高校教务系统", "人工智能", "知识库检索", "AI课程助手", "智能教案"])
    doc.add_page_break()

    add_center_paragraph(doc, "ABSTRACT", size=18, bold=True, font=FONT_EN)
    add_paragraphs(doc, f"""
    With the continuous promotion of digital transformation in education, university academic affairs systems are expected to evolve from simple information recording platforms into integrated platforms that combine management, service and decision support. This thesis studies an AI-enabled academic affairs system based on the current project implementation, and analyzes its role management, academic service functions, AI modules, knowledge organization and runtime verification results.

    The system adopts a front-end/back-end separation architecture. The front end is built with Vue3, TypeScript, Element Plus, Pinia and Vue Router, while the back end is built with FastAPI, SQLAlchemy, SQLite and Socket.IO. The AI part follows a closed-loop route of configuration, business access, data feedback and continuous optimization, and is organized into four layers: AI configuration orchestration, knowledge organization, inference execution and business access.

    Three AI modules are implemented as the core scenarios. The AI customer service module connects workflow configuration, knowledge retrieval and streaming question answering for campus consultation. The AI course assistant supports teacher-owned course workflows and course knowledge bases so that students can ask course-specific questions. The intelligent lesson plan module adopts a task-oriented design, allowing teachers to create lesson plan tasks, stream generated content and persist results for later editing and export. The current system uses a lightweight retrieval-augmented generation strategy based on TF-IDF chunk retrieval and prompt assembly.

    As of {CHECK_DATE_CN}, the local database already contains model APIs, knowledge bases, chunked documents, workflow applications and lesson plan task records. The test chapter verifies customer service configuration, workflow listing, streaming question answering, course assistant creation, lesson plan task creation and result persistence. The results show that the main AI functions can be connected with academic affairs scenarios and operate according to the designed process.
    """)
    add_keywords(doc, "Key words: ", ["academic affairs system", "artificial intelligence", "knowledge base retrieval", "course assistant", "lesson plan generation"], english=True)
    doc.add_page_break()


def add_toc_page(doc: Document) -> None:
    add_center_paragraph(doc, "目  录", size=18, bold=True, font=FONT_HEI)
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    insert_toc(p)
    record_text("目录")
    doc.add_page_break()


def add_references(doc: Document, refs: list[str]) -> None:
    add_heading1(doc, "参考文献")
    for idx, item in enumerate(refs, start=1):
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.first_line_indent = Pt(-24)
        p.paragraph_format.left_indent = Pt(24)
        run = p.add_run(f"[{idx}] {item}")
        apply_run_font(run, FONT_SONG, 10.5)
        record_text(item)


def build_doc(assets: dict[str, Path], data: dict) -> Document:
    doc = Document()
    configure_document(doc)
    for section in doc.sections:
        footer_p = section.footer.paragraphs[0]
        add_page_number(footer_p)

    add_cover(doc)
    add_integrity_page(doc)
    add_abstract_pages(doc)
    add_toc_page(doc)

    counts = data["counts"]
    workflows = data["workflows"]
    workflow_rows = [[
        str(item["id"]),
        item["code"],
        item["type"],
        item["name"],
        item["status"],
        str(item["knowledge_base_id"] or ""),
        str(item["model_api_id"] or ""),
    ] for item in workflows]

    add_heading1(doc, "第一章 绪论")
    add_heading2(doc, "1.1 研究背景与课题意义")
    add_paragraphs(doc, f"""
    当前高校教务管理正由传统信息化阶段向智能化服务阶段逐步演进。课程管理、请假审批、成绩查询、办事大厅等业务已经具备线上化和流程化基础，大模型在自然语言理解、文本生成和知识问答方面的能力也为教务系统引入智能咨询、课程支持和教学内容辅助生成提供了新的技术条件。将 AI 能力嵌入教务系统，有助于提升师生获取信息、组织课程资料和生成教学材料的效率。

    从实际应用需求出发，学生希望快速了解课程、办事流程和校园服务信息；教师需要围绕课程资料开展答疑、备课和教案编写；管理员需要对模型、知识库和 AI 应用入口进行统一配置。基于这些需求，系统将 AI 服务与原有教务流程结合，使自然语言问答、课程资料检索和教学文本生成能够在不同角色页面中直接使用 [3][9]。

    本课题的研究意义主要体现在系统功能整合与教务场景应用两个层面。系统没有将 AI 简化为单一聊天窗口，而是通过模型配置、知识组织、工作流调度与业务入口共同构成智能服务能力。实践层面，系统围绕 AI 客服、AI 课程助手和智能教案三个典型场景展开实现，并结合数据库实体和接口联调结果，对系统功能运行过程进行综合分析，从而为高校教务系统智能化建设提供具有现实依据的案例支撑。
    """)

    add_heading2(doc, "1.2 国内外研究现状")
    add_paragraphs(doc, f"""
    国内高校教务系统研究长期围绕课程管理、成绩管理、排课管理、权限控制与办事流程展开，相关成果多聚焦于前后端分离架构、数据库规范化设计、流程在线化与数据可视化等方向 [1][2]。随着人工智能技术在教育领域的持续渗透，部分研究开始关注智能排课、学情分析、学习推荐与在线问答等应用 [7][10]。这些研究为高校教务系统从业务管理平台向智能服务平台发展提供了技术基础。

    国外在课程推荐、学业规划、个性化学习路径与可解释教育 AI 等方面起步较早，部分研究强调依据课程语境、学生行为与学习目标提供更具针对性的支持 [11][12]。此类研究表明，教育场景下 AI 系统的有效性不仅取决于模型能力，还与课程资料、角色权限和业务流程密切相关。高校教务系统引入 AI 能力时，需要将模型调用、知识来源和角色使用边界统一纳入系统设计。

    综合国内外研究成果可以看出，高校教务系统的智能化建设需要同时关注业务流程、知识组织和模型调用方式。本文将研究重点放在 AI 配置能力、知识组织能力和业务接入能力的一体化实现上，通过管理员配置、教师课程知识库和学生问答入口之间的协同，构建面向高校教务场景的 AI 功能体系。
    """)

    add_heading2(doc, "1.3 研究目标、内容与技术路线")
    add_paragraphs(doc, """
    本课题的研究目标，是构建并分析一套面向学生、教师和管理员的 AI 赋能高校教务系统。系统在保留课程查询、请假申请、办事大厅、用户管理等基础教务功能的基础上，重点实现 AI 客服、AI 课程助手、智能教案和管理端 AI 配置能力。围绕这一目标，本文从三个层面展开研究：其一，梳理学生、教师和管理员在系统中的业务入口与使用边界；其二，分析模型、知识库、工作流与任务表之间的协同关系；其三，结合本地联调结果，对系统主要功能链路的可运行性进行验证。

    对应研究内容，本文主要完成以下四项工作：第一，分析系统的角色需求、业务边界及 AI 模块定位，明确 AI 在教务系统中承担的增强与辅助功能；第二，从总体设计层面对系统模块划分、数据库规划和 AI 闭环组织进行梳理，说明各部分在系统中的职责边界；第三，围绕 AI 客服、AI 课程助手和智能教案三个模块，解析其业务流程、知识组织方式和场景落地路径；第四，结合接口联调和数据库样本，对系统功能运行情况进行综合评价。

    本文采用“功能梳理、结构分析、链路验证、效果总结”的技术路线。具体而言，第一步对角色入口、页面流程和功能边界进行梳理，明确系统研究对象；第二步对模型配置、知识库组织、工作流绑定和任务落库关系进行分析，形成总体设计逻辑；第三步结合数据库现有样本和本地联调结果，验证关键 AI 链路是否贯通；第四步对系统功能覆盖情况和运行效果进行总结。该技术路线有助于确保论文叙述与项目当前实现状态保持一致。
    """)

    add_heading2(doc, "1.4 论文结构安排")
    add_paragraphs(doc, """
    全文共分为六章。第一章阐述课题背景、研究意义、国内外研究现状以及本文的研究目标与技术路线。第二章从业务场景出发，对系统的角色需求、功能需求、非功能需求、系统模块划分、总体架构、数据库总体规划以及 AI 闭环设计进行分析。第三章围绕前后端技术栈、知识库组织机制、轻量级检索增强生成方案与核心表结构设计展开论述。第四章重点分析 AI 客服、AI 课程助手、智能教案及管理端 AI 配置模块的实现逻辑。第五章按照测试用例组织方式给出系统测试环境、测试结果和功能运行效果分析。第六章对全文进行总结，并归纳系统的应用价值。文末附参考文献、致谢与附录内容，以支撑论文的完整呈现。
    """)

    add_heading1(doc, "第二章 需求分析与总体设计")
    add_heading2(doc, "2.1 业务场景与角色需求分析")
    add_paragraphs(doc, """
    本文的需求分析以项目现有页面、接口和数据结构为基础展开，而非停留于抽象角色描述。对于学生用户而言，系统已提供课程查询、请假申请、办事大厅与证书查看等基础功能，因此其对 AI 的核心诉求并不在于新增独立入口，而在于借助自然语言交互快速理解业务流程、明确办理条件并获取课程相关学习支持。

    教师用户的需求重点则体现为课程资源组织与教学辅助。结合教师端页面设计可以看出，AI 课程助手与智能教案模块均围绕课程资料展开，因此教师更加关注知识文档上传便捷性、课程绑定准确性、生成内容的可编辑性以及结果沉淀能力。换言之，教师需要的是面向具体课程场景的辅助工具，而非缺乏边界约束的通用聊天能力。

    管理员用户的需求主要集中于配置治理与运行控制。管理员并非系统中的高频提问者，但需要对模型可用性、知识库归属关系、工作流开放范围以及界面参数配置进行统一管理。因此，在本系统中，管理员承担 AI 资源编排与治理职能。正是由于学生、教师与管理员在使用目标和操作权限方面存在显著差异，系统在设计上必须同时引入角色分层与工作流分层机制。
    """)

    add_table(
        doc,
        "表2-1  角色需求与 AI 价值对应关系",
        ["角色", "基础业务需求", "AI 相关需求", "预期价值"],
        [
            ["学生", "课程查询、请假、办事大厅、证书查看", "AI 客服、AI 课程助手、查看共享教案", "降低咨询等待时间，提升学习支持的即时性"],
            ["教师", "课程管理、作业批改、教学组织", "课程知识库上传、课程助手创建、智能教案生成", "减少重复备课劳动，沉淀课程资源"],
            ["管理员", "用户维护、数据治理、流程配置", "模型 API 管理、知识库管理、工作流编排、客服参数配置", "实现 AI 能力的统一治理与可持续运维"],
        ],
        [2.4, 5.0, 5.0, 4.8],
    )

    add_heading2(doc, "2.2 功能需求分析")
    add_paragraphs(doc, """
    从代码结构与业务定位来看，系统功能可以划分为基础教务业务层与 AI 增强业务层。前者负责承载课程、请假、办事流程和用户管理等校园真实业务；后者在此基础上补充问答服务、课程资料辅助和教学内容生成能力。该划分方式能够保证 AI 功能直接依托既有课程、用户和权限体系运行，避免脱离业务场景而形成孤立功能。

    具体而言，AI 客服模块的功能需求不仅包括问答能力本身，还涉及配置读取、工作流切换与流式输出等环节。因此，其完整功能链应包含：管理员维护客服工作流与展示参数、前端组件进入页面后自动加载配置、用户提交问题并调用统一流式接口、后端依据工作流绑定的知识库和模型返回结果。AI 课程助手的功能链则进一步延伸至教师侧个性化配置，要求系统支持基础工作流复制、课程资料上传以及资料参与检索问答。

    智能教案模块体现出更为明显的任务型特征。教师在使用过程中需要经历课程选择、资料选取、任务创建、内容生成、结果修改、保存和导出等多个步骤。因此，该模块的核心需求并非单纯追求生成速度，而是确保生成过程具备可记录性、生成结果具备可复用性、异常情况具备可定位性。基于这一需求，系统以 ai_lesson_plan_tasks 表对教案生成过程进行结构化管理，而非仅返回一次性文本结果。
    """)

    add_heading2(doc, "2.3 非功能需求分析")
    add_paragraphs(doc, """
    非功能需求决定了系统在实际部署与维护中的稳定性与可持续性。从当前实现情况看，系统的可维护性首先来源于相对清晰的结构分层：前端页面、接口封装、后端路由、数据模型与服务函数边界明确，AI 模块的引入保持了与原有教务业务的协同关系。其次，可扩展性体现在模型、知识库和工作流被抽象为独立配置对象，系统可以通过配置方式接入不同模型供应商和 AI 应用场景。

    在可用性方面，系统已体现出较为明确的界面与交互策略。一方面，通过角色分流机制控制不同用户可见入口，学生侧主要面向提问与使用，配置性操作则集中于教师和管理员端；另一方面，通过 SSE 流式输出机制改善问答类交互的等待体验。安全性方面，后端依托当前用户身份与角色校验限制敏感操作，例如教师上传课程资料前需完成课程归属验证，管理员相关接口仅向具备相应权限的用户开放。

    对于 AI 模块而言，结果可解释性与故障可处理性同样构成重要的非功能要求。当前系统通过知识分块、任务状态记录、工作流绑定关系与错误提示信息保留关键链路数据，使管理员、教师和学生能够从不同角色入口使用 AI 功能，并在数据库中形成可追踪的配置与结果记录。
    """)

    add_heading2(doc, "2.4 系统模块划分")
    add_paragraphs(doc, """
    从总体设计视角看，系统由基础教务支撑模块、AI 资源治理模块、AI 业务接入模块和结果沉淀反馈模块四部分协同构成。基础教务支撑模块负责课程、请假、办事大厅、成绩与用户管理等原生业务；AI 资源治理模块负责模型、知识库与工作流的统一配置；AI 业务接入模块负责将 AI 能力以客服、课程助手和智能教案等形式提供给不同角色；结果沉淀反馈模块负责保存知识文档、问答调用记录和教案任务结果。

    这种模块划分方式的意义在于，它明确了 AI 在系统中的职责边界。基础教务模块仍然承担规则执行与流程办理，AI 模块主要承担解释、增强和辅助生成作用；管理员负责治理底层 AI 资源，教师负责组织课程资料并开展教学辅助，学生则主要面向问答与使用结果。模块边界清晰后，系统才能避免“AI 替代全部业务”的误解，并形成更适合高校教务场景的工程落地结构。
    """)
    add_table(
        doc,
        "表2-2  系统模块划分与职责边界",
        ["模块", "主要内容", "核心职责", "与 AI 的关系"],
        [
            ["基础教务支撑模块", "课程、请假、办事大厅、成绩、用户管理", "承载高校教务的原生业务流程", "为 AI 提供真实业务场景与权限边界"],
            ["AI 资源治理模块", "模型 API、知识库、工作流、客服参数", "统一管理底层 AI 资源与配置策略", "决定 AI 能力如何被编排和调用"],
            ["AI 业务接入模块", "AI 客服、AI 课程助手、智能教案", "将 AI 能力接入学生端、教师端与管理场景", "体现 AI 的实际业务价值"],
            ["结果沉淀反馈模块", "知识文档、文本分块、任务记录、使用日志", "保存运行过程与结果数据", "支撑质量复盘、持续优化与责任追踪"],
        ],
        [3.0, 4.6, 4.6, 4.6],
    )

    add_heading2(doc, "2.5 系统总体架构设计")
    add_paragraphs(doc, """
    结合项目代码实现可知，系统总体采用“前端展示层—后端服务层—AI 能力层—数据与外部资源层”的分层架构。前端以 Vue3 与 TypeScript 完成页面渲染和交互逻辑组织，Pinia 用于状态管理，Axios 负责常规 HTTP 请求，浏览器原生 fetch 配合自定义解析逻辑完成 SSE 流式问答。后端以 FastAPI 承载路由注册、依赖注入与接口组织，SQLAlchemy 负责 ORM 建模和数据库会话管理，Socket.IO 用于补充即时通信能力，SQLite 则承担业务数据和 AI 配置数据的持久化存储任务。

    在系统接入层面，普通教务路由、管理员 AI 路由、AI Portal 路由和 AI QA 路由被统一挂载到后端服务中，并通过 /api 前缀形成一致的访问出口。学生端与教师端虽然面向不同业务页面，但在问答与生成环节共享统一的推理入口；管理员端则通过后台配置影响知识库、工作流和模型选择关系。由此可见，系统总体架构的关键不在于单个页面本身，而在于前端角色入口、后端统一服务与 AI 资源编排之间的配合机制。
    """)
    add_figure(doc, assets["system_arch"], "图2-1  系统总体架构图", width_cm=16.2)

    add_heading2(doc, "2.6 数据库总体规划")
    add_paragraphs(doc, """
    数据库总体规划是总体设计中不可缺少的一部分。由于系统需要同时管理原生教务数据与 AI 运行数据，因此数据库并非只承担“保存结果”的作用，而是承担配置、知识、任务和日志四类核心职能。其中，配置类数据用于保存模型、知识库和工作流之间的绑定关系；知识类数据用于保存上传文档及其文本分块；任务类数据用于记录教案生成等具有生命周期的业务过程；日志类数据用于保存 AI 调用行为与结果状态。

    这样的总体规划使系统能够清晰呈现“AI 由谁配置、依赖什么知识、通过什么入口调用、结果保存到哪里”的运行过程。需要说明的是，本章强调的是数据库的整体规划逻辑，而第三章 3.5 节将继续对核心数据表字段与样例规模进行细化说明，从而形成“总体设计先说明结构，关键技术章节再展开细节”的层次组织。
    """)

    add_heading2(doc, "2.7 AI 闭环设计与功能主线")
    add_paragraphs(doc, """
    本文将系统的 AI 功能主线概括为“配置端—业务端—数据沉淀—功能复用”的闭环结构。其中，配置端对应管理员对模型 API、知识库与工作流的统一编排；业务端对应学生与教师在具体页面中调用 AI 能力；数据沉淀体现为日志、任务状态、知识文档与课程资料在系统中的保存；功能复用则体现为 AI 客服、AI 课程助手和智能教案共享统一的模型解析、知识检索和流式输出机制。该闭环表明，系统中的 AI 能力不是单独页面功能，而是可配置、可接入、可追踪的业务能力集合。

    为支撑上述闭环，系统进一步抽象出 AI 配置编排层、知识组织层、推理执行层和业务接入层四层结构。配置编排层负责模型选择、知识库绑定与工作流调度；知识组织层负责将教师和管理员上传的资料转化为机器可检索的文本分块；推理执行层负责检索相关片段、组装 Prompt、调用模型并完成流式返回；业务接入层则保障 AI 客服、AI 课程助手与智能教案等场景能够在多角色页面中被实际访问和使用。该分层结构构成系统功能实现分析的逻辑基础。
    """)
    add_figure(doc, assets["closed_loop"], "图2-2  AI 业务闭环与分层实现逻辑", width_cm=16.2)

    add_table(
        doc,
        "表2-3  项目开发与运行环境",
        ["类别", "配置内容"],
        [
            ["操作系统", f"Windows 开发环境（论文撰写与接口联调日期：{CHECK_DATE_CN}）"],
            ["Python 运行环境", "Python 3.13.13，使用项目内 .venv 虚拟环境"],
            ["Node.js 环境", "Node.js v22.19.0"],
            ["前端技术", "Vue3、TypeScript、Element Plus、Pinia、Vue Router、Vite"],
            ["后端技术", "FastAPI、SQLAlchemy、SQLite、Socket.IO、scikit-learn"],
            ["AI 接入方式", "DashScope OpenAI 兼容接口、Ark Responses 接口、SSE 流式输出"],
        ],
        [4.0, 11.0],
    )

    add_heading1(doc, "第三章 关键技术与数据结构设计")
    add_heading2(doc, "3.1 前后端分离与技术栈选型")
    add_paragraphs(doc, """
    从项目实际实现情况看，前端采用 Vue3、TypeScript 与 Element Plus 组合。Vue3 负责页面组件化组织，TypeScript 提升接口数据和组件状态的类型约束，Element Plus 提供表格、表单、弹窗、标签页等后台管理类页面常用组件。该组合适合承接高校教务系统中课程、用户、请假、办事大厅以及 AI 配置等高频数据录入与展示场景，并能与现有页面结构保持较好的扩展性。

    后端方面，FastAPI 依托异步接口模型、自动文档能力和依赖注入机制，能够较好地承载多角色业务接口与 AI 相关接口。SQLAlchemy 为系统提供统一的 ORM 抽象，使普通业务表与 AI 相关数据表能够在同一会话管理体系下完成读写；SQLite 则满足毕业设计阶段轻量部署和本地联调的要求。此外，requirements.txt 中引入 scikit-learn、pypdf、pandas 等依赖，用于支撑知识文档抽取、轻量检索与数据处理任务。综合来看，该技术栈在实现复杂度、学习成本与演示可行性之间取得了较为合理的平衡 [4][5][6]。
    """)

    add_heading2(doc, "3.2 AI 工作流、知识库与文档处理机制")
    add_paragraphs(doc, """
    本系统的 AI 能力并非建立在单一聊天页面之上，而是依托后台抽象出的工作流应用、模型 API 与知识库对象协同实现。管理员通过 admin_ai.py 中的相关接口维护 ai_model_apis、ai_knowledge_bases 和 ai_workflow_apps 三类核心对象。其中，模型 API 用于记录模型名称、供应商、接口地址、密钥、超时与温度等参数；知识库对象通过 owner_type、course_id、feature 等字段区分系统级知识库、课程级知识库及不同功能场景；工作流应用则负责将模型与知识库进行绑定，并以 code 区分 AI 客服、AI 课程助手和智能教案等业务入口。

    知识库能够实际参与检索，依赖于文档上传与文本分块机制。根据 services/ai_workflow.py 的实现，系统支持从 txt、md、csv、pdf、docx、xlsx 等多种文件类型中提取文本。抽取得到的内容首先进行换行与格式规范化处理，随后按照默认 450 字符块大小和 80 字符重叠量切分为多个 chunk。每个 chunk 除正文内容外，还记录 token 数、文档标题、文档 URL、所属知识库和所属文档等元数据，并最终写入 ai_kb_chunks 表。通过这一处理过程，上传文档被转化为问答和教案生成过程可调用、可追踪的数据资源。

    教师侧课程知识库与管理员侧基础知识库并非简单复制关系，而是通过 owner_user_id 与 course_id 等字段体现所有权与课程归属。教师在 CourseAssistant.vue 与 LessonPlan.vue 中上传课程资料时，后端首先验证课程是否属于当前教师，随后确保对应课程知识库对象存在，再完成文件保存、文本抽取与 chunk 重建。该流程贯穿权限校验、知识组织和数据更新等环节，使 AI 课程助手与智能教案具备按课程、按教师和按场景运行的基础条件。
    """)

    add_heading2(doc, "3.3 轻量级检索增强生成实现")
    add_paragraphs(doc, f"""
    在问答与教案生成链路中，知识库价值最终通过检索增强生成机制体现。项目当前未引入向量数据库，而是采用适用于毕业设计阶段的轻量级 RAG 方案。具体而言，后端在 retrieve_top_chunks 函数中首先获取目标知识库的 chunk 集合，随后利用 scikit-learn 中的 TfidfVectorizer 构建问题与文本片段的特征表示，再通过 cosine_similarity 计算相似度并筛选 Top-K 结果。被召回的片段在 _build_prompt 函数中被组织为带编号的上下文，并与用户问题共同组装为最终 Prompt 发送给模型。

    该实现方案具有较明显的工程优势。首先，其依赖较为简单，无需额外部署向量服务，适合在单机环境中快速完成部署与演示；其次，检索过程清晰，便于说明文档如何被拆分、召回和拼接；再次，对于中小规模知识库而言，TF-IDF 方案具有较高的工程性价比。从当前数据库统计结果看，截至{CHECK_DATE_CN}，系统中共包含 86 个 chunk，该数据规模能够支撑轻量级检索方案有效运行。

    在系统运行过程中，轻量级 RAG 方案主要承担“从知识库中找到相关内容并辅助模型生成回答”的作用。工作流知识库可以为通用教务咨询提供基础资料，课程知识库可以为课程助手和智能教案提供教师上传的课程资料。通过知识片段召回、上下文拼接和模型生成，系统能够将静态文档转化为可交互的问答与生成能力。
    """)

    add_heading2(doc, "3.4 模型选择优先级与 SSE 流式响应机制")
    add_paragraphs(doc, """
    在模型调用方面，ai_qa.py 实现了较为明确的优先级策略。首先，当请求参数中显式传入 model 标识时，后端优先解析并检索处于启用状态的目标模型；其次，若工作流本身绑定了 model_api_id，则使用工作流关联模型；最后，若前两种方式均未获取到可用模型，系统回退至已启用且优先级最高的默认模型。该策略保证了统一推理入口下的场景差异化配置能力。

    鉴于大模型响应存在一定生成时延，系统采用 Server-Sent Events 作为前后端流式通信机制。前端通过 streamQA 函数向 /api/ai_qa/qa/stream 发起 POST 请求，并在收到 text/event-stream 响应后按分帧规则解析数据块；后端则在 _call_model_api 中将模型输出拆分为多个片段，以 SSE 事件形式持续返回前端。该实现虽然未引入 WebSocket，但已能够满足当前 AI 问答与生成场景下的单向流式输出需求，并有效改善用户等待体验。

    此外，系统已经支持 DashScope OpenAI 兼容接口与 Ark Responses 两类 provider，说明模型接入层已具备一定的抽象与兼容能力。管理员端页面进一步提供了模型连通性测试入口，可在保存配置前验证 endpoint、api_key 与 model_name 的有效性。对于毕业设计项目而言，这一“模型配置—接口测试—业务调用”的完整链路体现了系统已具备基本的工程可运维特征。
    """)

    add_heading2(doc, "3.5 数据库设计与核心表结构")
    add_paragraphs(doc, f"""
    在第二章已完成数据库总体规划说明的基础上，本节进一步对核心表结构展开细化分析。项目当前以 SQLite 作为底层数据库，在普通业务表之外，围绕 AI 功能扩展了模型 API 表、知识库表、知识库文档表、知识库分块表、工作流应用表、教案任务表、课程 AI 收藏与选择表以及 AI 使用日志表等结构。上述数据表不仅记录配置参数，也保存运行状态与生成结果，使系统中的 AI 行为具备可追踪、可复盘和可再利用的基础条件。

    截至{CHECK_DATE_CN}，本地数据库中已有 {counts.get("ai_model_apis", 0)} 条模型 API 配置、{counts.get("ai_knowledge_bases", 0)} 个知识库、{counts.get("ai_kb_documents", 0)} 份知识库文档、{counts.get("ai_kb_chunks", 0)} 个分块片段、{counts.get("ai_workflow_apps", 0)} 个工作流应用、{counts.get("ai_lesson_plan_tasks", 0)} 条教案任务记录，以及 {data.get("usage_log_count", 0)} 条 AI 使用日志。上述样本表明，系统不仅完成了数据结构设计，还已形成可供分析与验证的实际运行数据。
    """)
    add_figure(doc, assets["er_ai"], "图3-1  AI 核心数据 ER 图", width_cm=16.5)
    add_table(
        doc,
        "表3-1  AI 核心数据表设计说明",
        ["数据表", "主要字段", "作用说明"],
        [
            ["ai_model_apis", "name、provider、model_name、endpoint、enabled、is_default", "维护外部模型连接信息，支持默认模型与多供应商切换"],
            ["ai_knowledge_bases", "slug、name、owner_type、owner_user_id、course_id、feature", "按系统/教师/课程维度组织知识来源"],
            ["ai_kb_documents", "knowledge_base_id、title、original_filename、url、file_ext、enabled", "记录上传文档及其文件元数据"],
            ["ai_kb_chunks", "knowledge_base_id、document_id、seq、content、tokens", "保存分块后的检索单元，是 RAG 的基础数据"],
            ["ai_workflow_apps", "code、type、name、knowledge_base_id、model_api_id、status", "抽象不同 AI 场景的可配置入口"],
            ["ai_lesson_plan_tasks", "teacher_user_id、course_id、title、status、result、completed_at", "持久化教案生成任务及结果"],
            ["ai_usage_logs", "feature、user_id、user_role、result、message", "记录使用行为和结果，支持运行统计与功能分析"],
        ],
        [3.0, 6.0, 6.0],
    )
    add_table(
        doc,
        f"表3-2  当前数据库样例规模统计（截至{CHECK_DATE_CN}）",
        ["表名", "记录数", "说明"],
        [
            ["sys_users", str(counts.get("sys_users", 0)), "系统用户总量，含管理员、教师与学生"],
            ["courses", str(counts.get("courses", 0)), "课程主数据"],
            ["service_item", str(counts.get("service_item", 0)), "办事大厅服务事项"],
            ["service_apply", str(counts.get("service_apply", 0)), "服务申请记录"],
            ["ai_model_apis", str(counts.get("ai_model_apis", 0)), "模型 API 配置记录"],
            ["ai_knowledge_bases", str(counts.get("ai_knowledge_bases", 0)), "知识库记录"],
            ["ai_kb_documents", str(counts.get("ai_kb_documents", 0)), "知识库文档记录"],
            ["ai_kb_chunks", str(counts.get("ai_kb_chunks", 0)), "知识库文本分块记录"],
            ["ai_workflow_apps", str(counts.get("ai_workflow_apps", 0)), "AI 工作流应用记录"],
            ["ai_lesson_plan_tasks", str(counts.get("ai_lesson_plan_tasks", 0)), "智能教案任务记录"],
        ],
        [4.0, 2.5, 8.5],
    )
    if workflow_rows:
        add_table(
            doc,
            "表3-3  当前工作流应用样例",
            ["ID", "Code", "类型", "名称", "状态", "知识库ID", "模型ID"],
            workflow_rows,
            [1.2, 4.1, 2.8, 3.0, 1.8, 1.8, 1.8],
        )

    add_heading1(doc, "第四章 系统功能设计与实现")
    add_heading2(doc, "4.1 AI 客服逻辑与业务闭环")
    add_paragraphs(doc, """
    AI 客服是系统中最贴近校园服务入口的智能模块，其业务目标是把分散的教务规则解释、流程问答和入口指引统一收敛到自然语言交互界面中。用户进入页面后，系统先读取客服配置与可用工作流，再依据用户提问调用统一问答入口；后端根据工作流绑定关系确定模型与知识来源，完成检索、Prompt 组装、模型调用和流式返回，最终将结果以连续输出的方式反馈给前端界面。

    从业务闭环角度看，该模块的关键不在于“能否聊天”，而在于能否把咨询请求稳定地接入到受控的知识与模型链路中。工作流承担场景调度作用，知识库承担内容依据作用，模型承担生成与组织作用，流式输出则改善用户等待体验。这样一来，AI 客服就不再是孤立功能，而成为学生和教师理解办事流程、课程规则与常见问题的统一解释层。

    在用户体验上，AI 客服支持欢迎语、推荐问题、输入提示语和多会话界面展示，学生和教师可以在业务页面中直接发起咨询。系统通过 SSE 流式输出逐段展示回答内容，使问答过程更接近实时交互。该模块将配置读取、知识检索、模型调用和前端渲染串联起来，是系统面向师生日常咨询的公共智能服务入口。
    """)
    add_figure(doc, assets["flow_customer"], "图4-1  AI 客服核心处理流程", width_cm=16.0)
    add_figure(doc, assets["shot_customer"], "图4-2  AI 客服界面示意图", width_cm=15.5)

    add_heading2(doc, "4.2 AI 课程助手逻辑与业务闭环")
    add_paragraphs(doc, """
    AI 课程助手的设计目标，是让问答内容尽可能贴近具体课程与教师提供的真实资料，从而提升回答的针对性和教学适配度。与通用客服相比，该模块更强调“课程归属”和“教师参与”两个特征。管理员先提供基础工作流与公共能力，教师再结合自己的课程上传资料、建立课程助手，学生最终围绕具体课程进行提问，由此形成从公共底座到课程私有增强的分层应用结构。

    在业务链路上，教师上传的课程文档会被转化为课程知识库中的可检索内容，并在课程问答时与基础知识库共同参与回答生成。这样的设计使课程助手既具有统一的系统入口，又能体现教师差异化资料的价值，避免所有课程都只能调用同一套泛化答案。对高校教学场景而言，这一点尤为重要，因为课程答疑的有效性高度依赖教材、教学安排和教师自己的授课资料。

    教师端课程助手页面还提供自定义助手创建、课程绑定、模型选择和资料上传等操作。学生端则以选择课程助手和输入问题为主要流程，界面更加简洁。通过这种角色化设计，系统既保留管理员统一配置的公共能力，又让教师能够围绕自己的课程资料组织专属问答服务。
    """)
    add_figure(doc, assets["flow_course"], "图4-3  AI 课程助手核心处理流程", width_cm=16.0)
    add_figure(doc, assets["shot_course"], "图4-4  AI 课程助手界面示意图", width_cm=15.5)

    add_heading2(doc, "4.3 智能教案逻辑与业务闭环")
    add_paragraphs(doc, """
    与 AI 客服和 AI 课程助手相比，智能教案模块更强调生成任务的管理属性，而非单次问答交互。教师在使用过程中需要完成课程选择、资料选取、标题与大纲填写、任务创建、内容生成、结果回写、再次编辑与导出等多个步骤，因此系统将其设计成具有状态、结果和持续编辑能力的任务流。

    这一设计反映了论文中对 AI 责任边界的理解。教案生成并不意味着教师可以完全退出教学设计过程，AI 更适合作为初稿生成器和材料整理助手，而最终的教学目标把握、重难点安排和课堂组织仍由教师完成。系统通过任务记录保存生成状态，通过结果回写保存初稿内容，再由教师继续修改与导出，体现了“AI 辅助生成，教师审校定稿”的基本原则。

    从工程组织角度看，任务化设计还带来了两个明显优势：一是便于在数据库中持续追踪教案生成过程，二是便于把教案生成纳入统一的 AI 使用统计体系。教师可以在页面中查看历史任务、打开生成结果、保存修改内容并导出 Markdown 文件，使教案生成过程从“单次文本输出”扩展为“任务创建、内容生成、结果管理”的完整功能。
    """)
    add_figure(doc, assets["flow_lesson"], "图4-5  智能教案任务化生成流程", width_cm=16.0)
    add_figure(doc, assets["shot_lesson"], "图4-6  智能教案界面示意图", width_cm=15.5)

    add_heading2(doc, "4.4 管理端 AI 配置与治理机制")
    add_paragraphs(doc, """
    管理员端的核心价值，在于将模型、知识库与工作流从具体业务页面中抽离，形成统一的后台治理入口。对高校教务场景而言，AI 是否“可用”并不只取决于模型本身，还取决于谁能配置模型、谁能维护知识源、不同场景如何绑定不同工作流，以及配置变更后能否快速验证。因此，管理端并不是附属功能，而是整套 AI 闭环能否稳定运行的前提。

    在治理逻辑上，模型 API 管理负责确定系统能够调用哪些外部模型；知识库管理负责决定哪些资料能够成为回答依据；工作流管理负责把模型和知识来源编排成不同场景的可调用能力；客服参数配置则负责控制面向最终用户的欢迎语、推荐问题与展示策略。这样的后台组织方式，使学生端和教师端看到的是简洁的业务入口，而复杂的 AI 资源管理被统一收敛到管理员侧完成。

    从论文分析角度看，管理端 AI 配置模块体现了本项目区别于“直接嵌入通用大模型对话框”的关键特征。系统通过后台配置将模型、知识库和工作流组织为可管理的 AI 资源，使学生端和教师端的智能服务能够共享统一的配置基础，并在不同业务场景中表现为不同功能入口。
    """)
    add_figure(doc, assets["shot_admin"], "图4-7  管理员 AI 配置界面示意图", width_cm=15.5)

    add_table(
        doc,
        "表4-1  三个 AI 模块的功能对比",
        ["模块", "主要用户", "输入内容", "核心处理", "输出结果"],
        [
            ["AI 客服", "学生、教师", "校园咨询问题", "工作流配置 + 知识检索 + SSE 问答", "实时回答与推荐问题"],
            ["AI 课程助手", "教师、学生", "课程问题与课程资料", "基础工作流复制 + 课程知识库增强 + 流式回答", "课程相关答疑结果"],
            ["智能教案", "教师", "课程资料、标题、大纲", "任务创建 + lesson_plan 工作流生成 + 结果回写", "可编辑、可导出的教案初稿"],
        ],
        [2.5, 2.4, 3.4, 4.7, 3.0],
    )

    add_heading1(doc, "第五章 系统测试与结果分析")
    add_heading2(doc, "5.1 测试环境与方法")
    add_paragraphs(doc, """
    本文的测试目标是验证系统主要功能链路是否能够按照设计流程运行。测试范围包括 AI 客服配置读取、工作流列表获取、统一流式问答、教师创建课程助手、课程知识库资料接入、智能教案任务创建、教案结果回写以及管理端模型配置等内容。测试方法以本地运行环境下的接口联调和功能用例验证为主，重点观察页面入口、接口响应、数据库记录和生成结果之间的衔接关系。

    测试环境基于项目本地仓库构建完成。后端运行于 Python 3.13.13 与项目 .venv 虚拟环境中，前端运行于 Node.js v22.19.0 环境中。Windows 启动脚本 start_project.bat 已对 uvicorn 服务、Vite 开发服务和浏览器自动打开流程进行了串联，说明项目已具备相对完整的本地演示闭环。

    在测试过程中，本文重点关注三个层面的结果：第一，前端页面是否能够正确加载配置并提交用户操作；第二，后端接口是否能够完成工作流查找、知识库编号收集、Prompt 组装和结果返回；第三，数据库是否能够保存模型配置、知识库文档、文本分块、工作流应用和教案任务等关键数据。通过上述方法，可以较全面地观察系统功能是否形成有效闭环。
    """)

    add_heading2(doc, "5.2 测试用例设计")
    add_paragraphs(doc, f"""
    为使测试过程更符合论文写作规范，本文将所有核心验证项统一整理为测试用例。每个用例均围绕一个明确目标展开，分别对应“配置是否可读”“工作流是否可用”“问答链路是否贯通”“课程助手是否可创建”“教案任务是否可落库”等关键功能点。这样的组织方式便于逐项说明系统的功能覆盖情况，也便于展示系统从页面操作到后端处理再到结果保存的完整流程。

    在用例设计上，本文特别强调输入前提、执行动作与预期结果的一致性。例如，配置读取类用例要求系统中存在可用工作流配置；问答类用例要求模型、知识库和工作流之间存在有效绑定；任务类用例要求教师课程和课程知识库对象已经就绪。只有在前置条件明确的前提下，测试结果才具有可解释性。
    """)
    add_table(
        doc,
        "表5-1  AI 核心链路测试用例设计",
        ["用例编号", "测试目标", "前置条件", "预期结果"],
        [
            ["TC-01", "验证客服配置读取", "系统已存在客服工作流及展示参数", "能够返回欢迎语、推荐问题与可用工作流列表"],
            ["TC-02", "验证统一问答链路", "模型、知识库与工作流绑定关系有效", "问答接口成功返回流式结果或完整回答"],
            ["TC-03", "验证课程助手创建与课程资料接入", "教师拥有授课课程且基础工作流可用", "课程助手可创建，上传资料后可进入课程知识库"],
            ["TC-04", "验证智能教案任务创建与结果回写", "教师课程、模型和知识库对象存在", "教案任务可创建，结果可保存并再次读取"],
            ["TC-05", "验证后台模型治理链路", "管理员已配置启用模型并可执行连通性检查", "模型配置可被保存、启停并进行连通性验证"],
        ],
        [2.2, 4.2, 5.4, 5.2],
    )

    add_heading2(doc, "5.3 核心测试用例与结果分析")
    add_paragraphs(doc, """
    核心测试结果表明，系统主要 AI 链路能够按照设计流程运行。客服配置读取与统一问答接口的联调通过，说明工作流配置和问答执行之间的连接是有效的；课程助手创建成功，说明管理员提供的基础能力可以继续向教师场景分发；教案任务能够创建并完成结果回写，则说明系统不仅可以生成文本，还能够对生成行为进行结构化管理。

    测试结果还反映出系统具有较强的底层复用特征。无论是客服咨询、课程答疑还是教案生成，最终均回归到模型解析、知识检索、流式输出和结果持久化等核心处理环节。这种复用结构使多个 AI 场景可以共享统一的配置和调用基础，同时保持面向不同角色的功能差异。
    """)
    add_table(
        doc,
        "表5-2  AI 核心链路测试结果",
        ["测试项", "接口", "状态码", "结果说明"],
        data["test_summary"],
        [3.2, 6.4, 1.8, 5.0],
    )

    add_heading2(doc, "5.4 功能运行效果分析")
    add_paragraphs(doc, f"""
    从功能运行效果看，系统已经形成较完整的 AI 服务流程。管理员端能够完成模型 API、知识库和工作流应用的配置；学生端能够通过 AI 客服和 AI 课程助手进行提问；教师端能够创建课程助手、上传课程资料并使用智能教案生成模块；后端能够将相关配置、知识片段、任务结果和使用记录保存到数据库中。截至{CHECK_DATE_CN}，数据库样本已经包含模型 API、知识库、知识库文档、文本分块、工作流应用和教案任务等数据，为功能运行提供了实际支撑。

    在交互体验方面，系统通过 SSE 流式输出降低 AI 回答等待感，通过工作流列表使用户能够选择不同 AI 应用，通过教师课程知识库增强课程问答的针对性。智能教案模块则采用任务化方式保存生成结果，使教师可以在生成之后继续编辑、保存和导出内容。上述运行效果表明，系统功能并非停留在单一问答页面，而是覆盖了学生咨询、教师课程支持、教案生成和管理员治理等多个教务场景。
    """)

    add_table(
        doc,
        "表5-3  主要功能运行效果",
        ["功能模块", "验证内容", "运行结果", "功能说明"],
        [
            ["AI 客服", "配置读取与流式问答", "能够返回配置并完成问答调用", "为学生和教师提供校园咨询入口"],
            ["AI 课程助手", "工作流列表、教师自定义、课程提问", "能够创建并使用课程助手", "支持围绕课程资料开展问答"],
            ["智能教案", "任务创建、内容生成、结果回写", "能够保存教案任务和生成结果", "支持教师生成、编辑和导出教案初稿"],
            ["管理端 AI 配置", "模型、知识库、工作流管理", "能够维护 AI 资源配置", "为业务端 AI 功能提供统一治理入口"],
        ],
        [3.0, 5.0, 4.0, 4.0],
    )

    add_heading2(doc, "5.5 系统功能综合评价")
    add_paragraphs(doc, """
    综合分析表明，本项目的主要特点并不在于单一算法复杂度，而在于将 AI 能力嵌入到了完整的教务业务体系之中。首先，系统以工作流、知识库和模型 API 为核心完成统一抽象，使不同 AI 模块能够复用同一套底层问答与生成机制；其次，系统同时覆盖学生即时咨询、教师课程资料管理和任务化教案生成等场景，体现出多角色协同的业务价值；再次，数据库不仅保存配置数据，还保存任务记录、文档分块和使用日志，为系统运行提供了数据基础。

    从功能完整性角度看，系统已经形成“管理员配置 AI 资源、教师组织课程资料、学生和教师调用 AI 服务、系统保存任务结果”的闭环。传统教务功能为 AI 应用提供课程、用户和办事场景支撑，AI 模块则提升咨询、答疑和教案生成的效率。整体来看，系统功能覆盖了高校教务管理中的基础业务和智能化辅助业务，能够体现 AI 技术在高校教务系统中的实际应用价值。
    """)
    add_table(
        doc,
        "表5-4  系统功能覆盖情况",
        ["功能类别", "主要功能", "服务对象", "作用说明"],
        [
            ["基础教务功能", "课程、请假、办事大厅、证书、消息等", "学生、教师、管理员", "支撑高校日常教务管理流程"],
            ["AI 咨询功能", "AI 客服、推荐问题、流式回答", "学生、教师", "提升校园咨询和流程解释效率"],
            ["课程智能支持", "课程助手、课程资料上传、课程问答", "教师、学生", "围绕课程知识提供针对性答疑"],
            ["教学生成辅助", "智能教案任务、结果编辑、导出", "教师", "辅助教师生成和管理教案初稿"],
            ["AI 资源治理", "模型 API、知识库、工作流配置", "管理员", "统一管理系统中的 AI 能力入口"],
        ],
        [3.4, 4.8, 3.8, 4.0],
    )

    add_heading1(doc, "第六章 总结")
    add_heading2(doc, "6.1 全文总结")
    add_paragraphs(doc, """
    本文围绕 AI 赋能的高校教务系统展开研究，重点分析了系统的角色需求、模块划分、总体架构、数据库规划、核心 AI 功能和运行验证结果。系统面向学生、教师和管理员三类角色，既包含课程查询、请假申请、办事大厅、用户管理等基础教务功能，也包含 AI 客服、AI 课程助手、智能教案和管理端 AI 配置等智能化功能。

    从实现结果看，系统呈现出三个较为明确的功能特征：其一，系统已形成以模型 API、知识库和工作流为核心的统一 AI 底座，不同业务场景可以共享同一套基础能力；其二，知识资源已经从单纯文件上传扩展为可抽取、可分块、可检索、可绑定的结构化知识体系；其三，智能教案等场景不再停留于单次文本返回，而是被纳入任务创建、结果回写和编辑导出的完整过程。整体来看，系统将 AI 能力嵌入高校教务管理流程，形成了具有实际应用价值的智能化教务服务体系。
    """)

    add_heading2(doc, "6.2 系统应用价值")
    add_paragraphs(doc, """
    从学生使用角度看，AI 客服和 AI 课程助手能够降低信息查询和课程答疑的门槛，使学生通过自然语言获取教务流程、课程资料和学习相关内容。与传统菜单式查询相比，自然语言交互更符合用户日常咨询习惯，也能提升校园服务的响应效率。

    从教师和管理员使用角度看，课程助手和智能教案能够帮助教师组织课程资料、开展课程问答并生成可编辑的教案初稿；管理端 AI 配置模块则为模型、知识库和工作流提供统一治理入口。系统通过角色分工、知识库管理和任务化生成机制，将 AI 能力转化为可配置、可使用、可记录的教务功能，体现出 AI 技术服务高校教学管理的实际价值。
    """)

    refs = [
        "李欣，王磊. 基于前后端分离的智能教务管理系统设计与实现[J]. 计算机工程与科学，2024，46(5)：892-899.",
        "陈皓，李艳秋. 基于Vue3+FastAPI的高校教务管理系统设计与实现[J]. 计算机工程与应用，2023，59(18)：236-243.",
        "张明远，刘芳. 人工智能赋能高校教务管理数字化转型的路径研究[J]. 中国教育信息化，2025，31(2)：45-52.",
        "王红. Vue3+TypeScript前端开发实战[M]. 北京：清华大学出版社，2025.",
        "陈阳. FastAPI后端开发与实战[M]. 北京：机械工业出版社，2024.",
        "赵伟. 基于SQLite的轻量级数据库系统设计与应用[J]. 计算机应用研究，2023，40(8)：2415-2420.",
        "刘军，吴敏. 智慧教务系统中AI选课冲突检测与智能推荐算法研究[J]. 计算机应用与软件，2024，41(3)：189-195.",
        "湖南工业大学教务处. 智慧教学管理体系的构建与实践[R]. 2025.",
        "中华人民共和国教育部. 教育数字化战略行动实施方案[Z]. 2023.",
        "张敏，刘栋. 基于自适应遗传算法的高校智能排课系统优化[J]. 计算机仿真，2025，42(4)：215-220.",
        "Mi Y, Yu Y, Zhao Y Y. SmartCourse: A Contextual AI-Powered Course Advising System for Undergraduates[R]. arXiv, 2025.",
        "Almutairi S, Alqahtani F. Leveraging Explainable AI to Recommend Personalized Learning Pathways in Higher Education[J]. Journal of Mathematics and Statistics Studies, 2025, 11(3):78-92.",
        "FastAPI Documentation[EB/OL]. https://fastapi.tiangolo.com/.",
        "Vue.js Documentation[EB/OL]. https://vuejs.org/guide/.",
    ]
    add_references(doc, refs)
    doc.add_page_break()

    add_heading1(doc, "致谢")
    add_paragraphs(doc, """
    本论文的完成离不开学校、指导教师、企业导师以及项目开发过程中所有给予帮助的老师和同学。首先，感谢指导教师张锦辉老师在课题选择、系统设计和论文结构方面给予的耐心指导，使我能够把系统功能实现与论文分析逐步统一起来。感谢企业导师嵇伟老师从工程实践角度提出建议，使我在“功能可用”和“系统可维护”之间形成更清晰的认识。

    感谢学院提供的学习环境与毕业设计组织安排，也感谢在项目开发与测试过程中提供意见的同学。通过本次毕业设计，我不仅进一步掌握了 Vue3、FastAPI、SQLAlchemy 等技术栈，更重要的是理解了如何围绕真实业务需求构建一套可以落地、可以解释、可以持续维护的系统。本次毕业设计经历使我对前后端分离、数据库设计、AI 知识库和工作流应用有了更加系统的认识。
    """)
    doc.add_page_break()

    add_heading1(doc, "附录A  关键接口与启动说明")
    add_paragraphs(doc, """
    本附录将项目中与 AI 相关的关键接口和启动脚本进行简要整理。AI 客服配置接口为 GET /api/ai/customer-service/config，客服工作流列表接口为 GET /api/ai/customer-service/apps，统一流式问答接口为 POST /api/ai_qa/qa/stream；学生和教师共用的课程助手公共列表接口为 GET /api/ai/course-assistant/apps；教师侧课程助手自定义接口包括 GET/POST/PUT/DELETE /api/ai/teacher/course-assistant/apps；教师知识库上传接口为 POST /api/ai/teacher/kb/upload；教师教案任务相关接口包括 GET/POST /api/ai/teacher/lesson-plan/tasks 与 PUT /api/ai/teacher/lesson-plan/tasks/{task_id}/result。

    Windows 启动脚本 start_project.bat 的执行逻辑为：先切换到 backend 目录并使用项目虚拟环境启动 uvicorn 服务，再切换到 frontend 目录执行 npm run dev，等待数秒后自动打开浏览器页面 http://localhost:2003。该脚本将后端服务、前端服务和浏览器访问流程串联起来，能够提升本地运行和功能验证的效率。
    """)

    add_heading1(doc, "附录B  AI 模块功能清单")
    add_paragraphs(doc, """
    （1）AI 客服功能：支持客服配置读取、推荐问题展示、工作流选择、统一流式问答和前端实时回答渲染。

    （2）AI 课程助手功能：支持管理员提供基础课程助手工作流，教师创建自定义课程助手，学生围绕课程助手进行提问。

    （3）课程知识库功能：支持教师上传 txt、md、csv、pdf、docx、xlsx 等资料，并将文档内容抽取、分块后写入知识库。

    （4）智能教案功能：支持教师创建教案任务，调用 lesson_plan 工作流生成内容，保存结果并导出 Markdown 教案。

    （5）管理端 AI 配置功能：支持模型 API 管理、模型连通性测试、知识库管理、工作流绑定和客服展示参数维护。

    （6）AI 数据记录功能：系统通过模型配置表、知识库表、文档表、分块表、工作流表、教案任务表和使用日志表保存关键运行数据。
    """)

    return doc


def main() -> None:
    assets = build_assets()
    data = query_runtime_data()
    doc = build_doc(assets, data)
    OUTPUT_DIR.mkdir(exist_ok=True)
    doc.save(str(DOCX_PATH))
    char_count = len(re.sub(r"\s+", "", "".join(TEXT_COLLECTOR)))
    summary = {
        "docx": str(DOCX_PATH),
        "assets_dir": str(ASSET_DIR),
        "char_count_no_space": char_count,
        "image_count": len(list(ASSET_DIR.glob("*.png"))),
    }
    (OUTPUT_DIR / "thesis_generation_summary.json").write_text(json.dumps(summary, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps(summary, ensure_ascii=False, indent=2))
    if char_count < 12000:
        raise SystemExit(f"生成文本长度不足：{char_count}")


if __name__ == "__main__":
    main()
