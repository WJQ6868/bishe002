from __future__ import annotations

import shutil
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm


DESKTOP = Path.home() / "Desktop"
DOC_PATH = next(DESKTOP.glob("*0518*副本.docx"))
BACKUP_PATH = DOC_PATH.with_name(
    DOC_PATH.stem + f"-backup-before-ch3-{datetime.now().strftime('%Y%m%d-%H%M%S')}" + DOC_PATH.suffix
)
ER_IMAGE = Path(r"D:\bishe\one\generated\thesis_assets\fig_3_1_ai_er_readable.png")


def find_paragraph(doc: Document, candidates: list[str]):
    for p in doc.paragraphs:
        text = p.text.strip()
        if text in candidates:
            return p
    raise ValueError(f"Paragraph not found: {candidates}")


def find_paragraph_startswith(doc: Document, prefixes: list[str]):
    for p in doc.paragraphs:
        text = p.text.strip()
        if any(text.startswith(prefix) for prefix in prefixes):
            return p
    raise ValueError(f"Paragraph not found by prefix: {prefixes}")


def set_paragraph_text(paragraph, text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def insert_normal_before(paragraph, text: str):
    new_p = paragraph.insert_paragraph_before(text)
    new_p.style = "Normal"
    return new_p


def insert_caption_before(paragraph, text: str):
    new_p = paragraph.insert_paragraph_before(text)
    new_p.style = "Caption"
    return new_p


def has_exact_paragraph(doc: Document, text: str) -> bool:
    return any(p.text.strip() == text for p in doc.paragraphs)


def main() -> None:
    shutil.copy2(DOC_PATH, BACKUP_PATH)
    doc = Document(str(DOC_PATH))

    # 3.1 系统总体功能设计
    set_paragraph_text(
        find_paragraph_startswith(
            doc,
            ["本章在前述需求分析基础上，对系统整体设计进行说明。"],
        ),
        "本章在前述需求分析基础上，对系统整体设计展开说明。由于本系统的研究重点在于为现有教务系统引入AI增强能力，因此本章不再重复铺陈全部基础业务页面的实现细节，而是重点说明基础教务业务、AI服务链路、数据资源与外部模型之间如何形成稳定的协同关系。换言之，第三章关注的是“系统应如何组织”，第四章再进一步说明“这些设计如何落到实际页面与代码之中”。",
    )
    set_paragraph_text(
        find_paragraph_startswith(
            doc,
            ["从系统设计角度看，基础教务功能与AI功能并非相互分离"],
        ),
        "从总体结构看，系统采用“基础教务业务层+AI增强服务层”的设计思路。基础教务业务层继续承担用户、课程、选课、成绩、请假、作业、办事和日程等确定性业务，保证规则执行、流程办理和权限控制的稳定性；AI增强服务层则围绕咨询解释、课程答疑、资料组织和教案生成提供补充能力。这样的划分有两层意义：一方面保留传统教务系统流程清晰、边界明确的特点，另一方面避免AI模块直接介入业务结果修改，从而使智能能力更适合承担解释、辅助和生成职责。",
    )
    set_paragraph_text(
        find_paragraph_startswith(
            doc,
            ["从设计上看，在现有教务系统中接入AI能力"],
        ),
        "在总体设计层面，学生、教师和管理员分别从不同业务入口接入系统，前端页面、后端服务、AI编排能力和数据资源在内部形成分层协同关系。图3-1主要用于说明这一总体结构：它强调AI能力为何能够嵌入现有教务系统，而不是作为独立工具悬浮在系统之外；同时也说明基础教务业务为何仍然是系统运行的主体，而AI模块更多承担增强和支撑作用。",
    )
    set_paragraph_text(
        find_paragraph(
            doc,
            ["图2-2基于现有教务系统的AI增强总体架构图", "图3-1 基于现有教务系统的AI增强总体架构图"],
        ),
        "图3-1 基于现有教务系统的AI增强总体架构图",
    )

    # 3.2 AI模块分层与协同设计
    set_paragraph_text(
        find_paragraph(doc, ["AI模块设计", "AI模块分层与协同设计"]),
        "AI模块分层与协同设计",
    )
    set_paragraph_text(
        find_paragraph_startswith(
            doc,
            ["从系统设计角度看，AI能力不应作为脱离教务业务的独立模块存在"],
        ),
        "在总体结构明确之后，还需要进一步说明AI模块内部如何分工。这里所说的“模块”，并不是单指某一个聊天页面，而是指从用户发起请求到系统返回结果这一过程中涉及的多类设计要素，包括角色入口、交互承接页面、接口封装、后端路由、工作流对象以及核心数据实体。将这些要素分层梳理，有助于解释AI客服、AI课程助手和智能教案虽然面向不同场景，却能够复用同一套底层AI资源与调用链路，见图3-2。",
    )
    set_paragraph_text(
        find_paragraph(
            doc,
            ["图3-1AI模块架构、代码组织与接口关系图", "图3-2 AI模块分层与接口协同关系图"],
        ),
        "图3-2 AI模块分层与接口协同关系图",
    )
    set_paragraph_text(
        find_paragraph_startswith(
            doc,
            ["图3-1将AI模块从角色入口到数据表的关系划分为五个层次"],
        ),
        "图3-2采用由外到内的方式描述AI模块的协同关系。最外层是角色入口层，用于说明学生、教师和管理员分别从咨询、课程学习、备课和配置治理等场景进入系统；其内侧是交互承接层，主要负责接收输入、展示状态和返回结果；再向内是接口收口层，用于把不同页面请求统一汇入后端服务；之后是AI能力编排层，负责按场景选择工作流、模型和知识来源；最内层则是数据与资源层，用于保存模型配置、知识库、文档片段、任务结果和使用日志。这一分层方式强调的是职责边界，而不是在总体设计图中堆叠全部实现细节。",
    )

    table_31_caption = find_paragraph(
        doc,
        ["表3-1AI功能组件、接口与核心类关系", "表3-1 AI功能组件、接口与核心类关系"],
    )
    set_paragraph_text(table_31_caption, "表3-1 AI功能组件、接口与核心类关系")

    if not has_exact_paragraph(
        doc,
        "为使图3-2中的设计要素能够与后续章节一一对应，表3-1进一步列出主要前端组件、接口封装、关键接口以及后端路由与核心类之间的关系。其中，前端组件是指直接承接用户交互的页面或组件；接口封装是指对前端HTTP请求进行统一包装的脚本；关键接口是指场景实际调用的API入口；后端路由与核心类则分别对应FastAPI中的接口入口与SQLAlchemy数据实体。",
    ):
        insert_normal_before(
            table_31_caption,
            "为使图3-2中的设计要素能够与后续章节一一对应，表3-1进一步列出主要前端组件、接口封装、关键接口以及后端路由与核心类之间的关系。其中，前端组件是指直接承接用户交互的页面或组件；接口封装是指对前端HTTP请求进行统一包装的脚本；关键接口是指场景实际调用的API入口；后端路由与核心类则分别对应FastAPI中的接口入口与SQLAlchemy数据实体。",
        )
        insert_normal_before(
            table_31_caption,
            "在这一层不再把所有字段、算法和页面状态全部压入总体设计图，而是将能够支撑设计说明的关键承载点单独整理出来。这样处理的目的，是把“系统总体结构”与“模块内部协同关系”区分开来，避免第三章在总体设计部分过早落入实现细节。",
        )

    table_31 = doc.tables[5]
    table_31.cell(1, 3).text = "ai_portal.py；ai_qa.py；AiWorkflowApp；AiModelApi；AiKnowledgeBase；AiKnowledgeBaseChunk"
    table_31.cell(1, 4).text = "面向通用教务咨询场景，统一读取配置并通过问答链路流式返回解释结果"
    table_31.cell(2, 3).text = "ai_portal.py；ai_qa.py；AiWorkflowApp；AiKnowledgeBase；AiKnowledgeBaseDocument；AiKnowledgeBaseChunk"
    table_31.cell(2, 4).text = "围绕课程与教师资料组织专属问答场景，支撑定向答疑"
    table_31.cell(3, 4).text = "先形成任务记录，再完成流式生成、结果回写与后续编辑导出"
    table_31.cell(4, 4).text = "统一维护模型、知识库、工作流与客服展示参数，为前台调用提供治理底座"

    heading_325 = find_paragraph(doc, ["AI工作流与知识资源组织设计"])
    if not has_exact_paragraph(
        doc,
        "从表3-1可以看出，三个AI业务入口虽然面向的用户和界面形态不同，但在设计上都经过统一接口收口，并共享模型配置、知识组织和工作流调度机制。这样一来，系统后续若继续扩展新的AI场景，不必重复建设底层能力，只需在既有设计框架上补充新的入口与场景配置即可。",
    ):
        heading_325.insert_paragraph_before(
            "从表3-1可以看出，三个AI业务入口虽然面向的用户和界面形态不同，但在设计上都经过统一接口收口，并共享模型配置、知识组织和工作流调度机制。这样一来，系统后续若继续扩展新的AI场景，不必重复建设底层能力，只需在既有设计框架上补充新的入口与场景配置即可。"
        ).style = "Normal"

    # 3.3 AI工作流与知识资源组织设计
    set_paragraph_text(
        find_paragraph_startswith(
            doc,
            ["AI功能没有被设计成单一的泛化聊天框"],
        ),
        "工作流（workflow）在本系统中并不是简单的流程图概念，而是对“场景—模型—知识来源”绑定关系的抽象。管理员可分别为AI客服、AI课程助手和智能教案建立不同工作流，使同一套模型接入能力能够根据场景切换不同知识库、模型参数和展示配置。这样的设计避免了把模型参数直接散落在各个页面中，也使启停控制、统一治理和后续扩展更为方便。",
    )
    set_paragraph_text(
        find_paragraph_startswith(
            doc,
            ["在知识来源方面，系统将其划分为公共知识和课程私有知识两部分"],
        ),
        "知识资源组织则采用“知识库—文档—文本片段”三级结构。知识库用于区分制度说明、课程资料等不同来源边界；文档用于保存上传资料及其文件元数据；文本片段则是文档经过抽取、清洗和切分后形成的最小检索单元。公共知识库主要服务于AI客服等通用咨询场景，课程知识库则按照教师、课程与资料归属组织，用于支撑课程助手与智能教案。",
    )
    set_paragraph_text(
        find_paragraph_startswith(
            doc,
            ["在文档处理设计上，资料上传并不代表流程结束"],
        ),
        "文档处理设计强调“上传—抽取—规范化—分块—入库”的连续过程。结合当前项目代码，系统支持对PDF、DOCX、TXT等资料进行文本抽取，并将内容按固定片段长度切分，同时保留一定重叠区间，以降低知识点恰好落在分块边界时造成的语义断裂。每个片段同时记录文档标题、来源地址和顺序编号，目的是使后续检索、定位与追踪具有明确依据。",
    )

    # 3.4 轻量级检索增强生成方案设计
    set_paragraph_text(
        find_paragraph_startswith(
            doc,
            ["在本项目里，轻量级检索增强生成并不是为了把算法写得复杂"],
        ),
        "检索增强生成（Retrieval-Augmented Generation，RAG）是指系统在调用大模型生成回答之前，先从本地知识库中召回相关内容，再将问题与召回结果共同送入模型。本系统采用这一方案，核心原因在于教务问答和课程答疑不能完全依赖通用知识自由生成，而应尽量建立在校内制度、课程资料和教师上传材料基础上，从而降低回答偏离业务规则的风险。",
    )
    set_paragraph_text(
        find_paragraph_startswith(
            doc,
            ["没有再引入更重的向量数据库"],
        ),
        "在检索层，系统采用TF-IDF（Term Frequency-Inverse Document Frequency，词频-逆文档频率）方法对问题与文本片段进行特征表示，并按照相似度排序筛选候选上下文。相较于额外部署向量数据库，这种设计更适合当前知识规模和毕业设计的部署条件：一方面依赖较少、便于本地与服务器环境统一维护，另一方面匹配过程更直观，更便于解释回答依据来自哪些文档片段。",
    )
    set_paragraph_text(
        find_paragraph_startswith(
            doc,
            ["从本课题的完成目标看，设计阶段更需要证明“资料能够被系统稳定调用”"],
        ),
        "结合当前项目实现，知识文档默认按约450个字符进行切分，并保留约80个字符的重叠区间；检索阶段再从候选片段中选取相关度较高的若干结果参与上下文组装。这样的参数设置并非单纯追求算法复杂度，而是在文档可读性、检索效率和部署成本之间取得平衡。对于高校教务规则、课程讲义和教师资料这类中小规模文本集合，这种轻量级方案已经能够满足问答与生成的基本需要。",
    )
    set_paragraph_text(
        find_paragraph_startswith(
            doc,
            ["在生成环节，系统不会把召回文本原样堆给模型"],
        ),
        "在生成层，系统并不会把召回文本原样返回，而是先依据场景确定知识范围，再将片段整理为上下文并与用户问题一并交由模型处理。这样一来，系统级知识和课程级知识能够形成分工：AI客服更偏向制度与流程解释，课程助手更贴近课程资料，智能教案则围绕课程主题和教学大纲生成结构化内容。由此，RAG方案成为连接知识组织层与模型调用层的关键设计环节。",
    )

    # 3.5 模型选择策略与SSE流式响应设计
    set_paragraph_text(
        find_paragraph_startswith(
            doc,
            ["模型调用层的设计重点不只是完成模型接入"],
        ),
        "模型选择策略的设计重点，不只是完成模型接入，还要保证不同场景下具有可切换能力。系统首先检查请求是否显式指定模型，其次检查当前工作流是否已经绑定模型；若前两步均未获取到可用配置，则回退至默认模型。这样的优先级顺序能够兼顾管理员统一治理与不同业务场景差异化配置两类需求，也使同一套问答入口在不同场景中保持灵活性。",
    )
    set_paragraph_text(
        find_paragraph_startswith(
            doc,
            ["响应方式上，系统采用Server-SentEvents作为流式输出机制"],
        ),
        "响应方式上，系统采用SSE（Server-Sent Events，服务器发送事件）作为流式输出机制。SSE适用于前端发起请求、后端持续回传结果的单向通信场景，与本系统中的问答和教案生成过程较为契合。采用这种方式后，用户在等待回答或教案内容生成时，可以持续看到输出变化，而不是长时间停留在无反馈状态，从而改善交互体验。",
    )
    set_paragraph_text(
        find_paragraph_startswith(
            doc,
            ["此外，模型接入层保留了provider级适配能力"],
        ),
        "此外，模型接入层保留了provider级适配能力，即在模型服务提供商发生变化时，系统不必整体重写业务逻辑，而只需围绕模型名称、接口地址、密钥、超时等参数进行配置调整。结合管理端提供的连通性校验功能，模型配置、接口测试和场景调用被组织为相对分离的设计层次，这为后续替换模型服务或补充新模型提供了扩展空间。",
    )

    # 3.6 数据库设计与核心表结构
    set_paragraph_text(
        find_paragraph_startswith(
            doc,
            ["数据库设计的重点不是单纯列出数据表"],
        ),
        "数据库设计采用“基础教务数据与AI扩展数据分层存储”的思路。基础教务数据继续承载用户、课程、选课、请假、办事和作业等确定性业务；AI扩展数据重点保存模型配置、知识资源、文档分块、工作流、教案任务和使用日志。这样设计的原因在于，AI能力并不是孤立存在，而是以基础教务数据给出的角色、课程和权限边界为前提，再以扩展数据完成知识组织和生成支撑。",
    )
    set_paragraph_text(
        find_paragraph_startswith(
            doc,
            ["进一步看，模型表解决“调用谁”的问题"],
        ),
        "从AI扩展数据之间的关系看，AiModelApi用于描述外部模型连接信息，可与多个AiWorkflowApp和AiLessonPlanTask形成关联；AiKnowledgeBase作为知识来源容器，可继续关联AiKnowledgeBaseDocument、AiKnowledgeBaseChunk、AiWorkflowApp和AiLessonPlanTask；AiKnowledgeBaseDocument记录原始资料，AiKnowledgeBaseChunk保存检索片段，两者之间形成一对多关系。上述关系构成了AI配置、知识调用和结果沉淀的核心链路，见图3-3；在此基础上，表3-2进一步列出各核心数据表的主要字段与作用说明。",
    )
    set_paragraph_text(
        find_paragraph(
            doc,
            ["表3-2AI核心数据表设计说明", "表3-2 AI核心数据表设计说明"],
        ),
        "表3-2 AI核心数据表设计说明",
    )

    table_32_caption = find_paragraph(doc, ["表3-2 AI核心数据表设计说明"])
    if not has_exact_paragraph(doc, "图3-3 AI核心数据ER图"):
        pic_para = table_32_caption.insert_paragraph_before()
        pic_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = pic_para.add_run()
        run.add_picture(str(ER_IMAGE), width=Cm(16.2))
        cap_para = insert_caption_before(table_32_caption, "图3-3 AI核心数据ER图")
        cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        insert_normal_before(
            table_32_caption,
            "图3-3所强调的不是所有字段细节，而是“模型—工作流—知识库—文档—分块—任务”之间的主干关系。系统在进行AI客服问答时，主要沿工作流、模型和知识库关系完成调用；在进行课程助手和智能教案处理时，则进一步叠加课程归属和任务记录关系。这样的关系设计使系统既能支持统一治理，又能保留课程场景下的差异化知识来源。",
        )

    table_32 = doc.tables[6]
    table_32.cell(4, 2).text = "保存分块后的检索单元，是轻量级RAG链路中的基础数据"
    table_32.cell(5, 2).text = "抽象不同AI场景的可配置入口，完成模型与知识库绑定"
    table_32.cell(6, 2).text = "持久化教案生成任务及结果，支撑状态追踪与结果回写"

    doc.save(str(DOC_PATH))
    print(f"Updated: {DOC_PATH}")
    print(f"Backup: {BACKUP_PATH}")


if __name__ == "__main__":
    main()
