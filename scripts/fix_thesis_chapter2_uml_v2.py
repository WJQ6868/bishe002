from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm
from docx.text.paragraph import Paragraph
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
ASSET_DIR = ROOT / "generated" / "thesis_assets"
FONT_REGULAR = Path(r"C:\Windows\Fonts\msyh.ttc")
FONT_BOLD = Path(r"C:\Windows\Fonts\simhei.ttf")


def get_input_doc() -> Path:
    desktop = Path.home() / "Desktop"
    matches = list(desktop.glob("2405273202*chapter2-updated.docx"))
    if not matches:
        raise FileNotFoundError("Cannot find chapter2-updated thesis file on Desktop.")
    return matches[0]


def get_output_doc(input_doc: Path) -> Path:
    return input_doc.with_name(input_doc.stem + "-v2.docx")


def ensure_backup(path: Path) -> Path:
    backup = path.with_name(path.stem + "-backup-before-uml-v2" + path.suffix)
    if not backup.exists():
        shutil.copy2(path, backup)
    return backup


def load_font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont:
    font_path = FONT_BOLD if bold and FONT_BOLD.exists() else FONT_REGULAR
    return ImageFont.truetype(str(font_path), size=size)


def wrap_text(draw: ImageDraw.ImageDraw, text: str, font: ImageFont.FreeTypeFont, max_width: int) -> list[str]:
    if not text:
        return [""]
    lines: list[str] = []
    current = ""
    for ch in text:
        candidate = current + ch
        if draw.textlength(candidate, font=font) <= max_width:
            current = candidate
        else:
            if current:
                lines.append(current)
            current = ch
    if current:
        lines.append(current)
    return lines or [text]


def draw_center_text(draw: ImageDraw.ImageDraw, center: tuple[int, int], text: str, font, fill="black") -> None:
    bbox = draw.textbbox((0, 0), text, font=font)
    w = bbox[2] - bbox[0]
    h = bbox[3] - bbox[1]
    draw.text((center[0] - w / 2, center[1] - h / 2), text, font=font, fill=fill)


def draw_box_text_centered(
    draw: ImageDraw.ImageDraw,
    box: tuple[int, int, int, int],
    title: str,
    lines: list[str],
    *,
    title_size: int,
    body_size: int,
    outline_width: int = 4,
    rounded: int = 18,
) -> None:
    x1, y1, x2, y2 = box
    draw.rounded_rectangle(box, radius=rounded, outline="black", fill="white", width=outline_width)

    title_font = load_font(title_size, bold=True)
    body_font = load_font(body_size)
    max_width = x2 - x1 - 30

    wrapped_title = wrap_text(draw, title, title_font, max_width)
    wrapped_body: list[str] = []
    for line in lines:
        wrapped_body.extend(wrap_text(draw, line, body_font, max_width))

    title_h = title_size + 8
    body_h = body_size + 10
    total_height = len(wrapped_title) * title_h + 12 + len(wrapped_body) * body_h
    start_y = y1 + max(18, (y2 - y1 - total_height) // 2)
    cx = (x1 + x2) // 2

    y = start_y
    for line in wrapped_title:
        draw_center_text(draw, (cx, y + title_size // 2), line, title_font)
        y += title_h
    y += 12
    for line in wrapped_body:
        draw_center_text(draw, (cx, y + body_size // 2), line, body_font)
        y += body_h


def draw_arrow(draw: ImageDraw.ImageDraw, start: tuple[int, int], end: tuple[int, int], width: int = 5) -> None:
    import math

    draw.line([start, end], fill="black", width=width)
    angle = math.atan2(end[1] - start[1], end[0] - start[0])
    arrow_len = 18
    angle1 = angle - math.pi / 8
    angle2 = angle + math.pi / 8
    p1 = (end[0] - arrow_len * math.cos(angle1), end[1] - arrow_len * math.sin(angle1))
    p2 = (end[0] - arrow_len * math.cos(angle2), end[1] - arrow_len * math.sin(angle2))
    draw.polygon([end, p1, p2], fill="black")


def draw_actor(draw: ImageDraw.ImageDraw, x: int, y: int, label: str) -> None:
    head_r = 24
    draw.ellipse((x - head_r, y, x + head_r, y + head_r * 2), outline="black", width=4)
    draw.line((x, y + 48, x, y + 132), fill="black", width=4)
    draw.line((x - 42, y + 76, x + 42, y + 76), fill="black", width=4)
    draw.line((x, y + 132, x - 36, y + 186), fill="black", width=4)
    draw.line((x, y + 132, x + 36, y + 186), fill="black", width=4)
    draw_center_text(draw, (x, y + 232), label, load_font(34, bold=True))


def draw_oval_usecase(draw: ImageDraw.ImageDraw, box: tuple[int, int, int, int], text: str) -> None:
    x1, y1, x2, y2 = box
    draw.ellipse(box, outline="black", fill="white", width=4)
    font = load_font(30, bold=False)
    lines = wrap_text(draw, text, font, x2 - x1 - 50)
    line_h = 40
    total_h = len(lines) * line_h
    y = y1 + (y2 - y1 - total_h) // 2 + line_h // 2 - 2
    cx = (x1 + x2) // 2
    for line in lines:
        draw_center_text(draw, (cx, y), line, font)
        y += line_h


def draw_label(draw: ImageDraw.ImageDraw, center: tuple[int, int], text: str, size: int = 24) -> None:
    draw_center_text(draw, center, text, load_font(size, bold=True))


def generate_usecase_figure(path: Path) -> None:
    img = Image.new("RGB", (2800, 1700), "white")
    draw = ImageDraw.Draw(img)
    draw_center_text(draw, (1400, 70), "系统角色需求用例图", load_font(54, bold=True))

    boundary = (530, 160, 2280, 1540)
    draw.rectangle(boundary, outline="black", width=5)
    draw.text((560, 180), "AI赋能高校教务系统", font=load_font(38, bold=True), fill="black")

    draw_actor(draw, 240, 320, "学生")
    draw_actor(draw, 240, 760, "教师")
    draw_actor(draw, 2580, 540, "管理员")

    usecases = {
        "课程查询与办理": (760, 330, 1220, 510),
        "AI客服咨询": (1380, 300, 1840, 500),
        "课程答疑支持": (1380, 580, 1840, 780),
        "课程资料上传与沉淀": (760, 850, 1320, 1070),
        "智能教案生成与修改": (1420, 930, 1980, 1150),
        "模型与知识库配置": (1420, 1240, 1980, 1440),
        "工作流编排与治理": (760, 1230, 1320, 1450),
    }
    for title, box in usecases.items():
        draw_oval_usecase(draw, box, title)

    lines = [
        ((240, 430), (760, 420)),
        ((240, 430), (1380, 400)),
        ((240, 430), (1380, 670)),
        ((240, 870), (1380, 670)),
        ((240, 870), (1040, 960)),
        ((240, 870), (1700, 1040)),
        ((2580, 650), (1700, 1340)),
        ((2580, 650), (1040, 1340)),
    ]
    for start, end in lines:
        draw.line([start, end], fill="black", width=4)

    draw_label(draw, (1230, 1520), "学生侧强调“能快速理解并继续办理”，教师侧强调“能围绕课程持续复用”，管理员侧强调“能统一配置并稳定治理”。", 22)
    img.save(path)


def generate_system_architecture_figure(path: Path) -> None:
    img = Image.new("RGB", (2600, 1900), "white")
    draw = ImageDraw.Draw(img)
    draw_center_text(draw, (1300, 70), "系统总体架构图", load_font(54, bold=True))

    top_boxes = [
        ((110, 170, 760, 430), "学生端", ["课程查询", "请假办理", "办事大厅", "AI客服", "AI课程助手"]),
        ((975, 170, 1625, 430), "教师端", ["课程管理", "成绩与作业处理", "课程资料上传", "AI课程助手", "智能教案"]),
        ((1840, 170, 2490, 430), "管理员端", ["用户维护", "模型接入管理", "知识库管理", "工作流编排", "客服参数维护"]),
    ]
    for box, title, lines in top_boxes:
        draw_box_text_centered(draw, box, title, lines, title_size=40, body_size=28)

    layers = [
        ((220, 580, 2380, 790), "前端展示层", ["基于角色展示业务入口、问答入口与任务入口，并负责状态反馈和结果呈现。"]),
        ((220, 900, 2380, 1110), "后端服务层", ["统一处理权限校验、业务编排、接口组织与普通教务业务和AI业务之间的衔接。"]),
        ((220, 1220, 2380, 1430), "AI能力层", ["完成模型调用、知识匹配、上下文组织、问答生成和教案生成等核心处理。"]),
        ((220, 1540, 2380, 1750), "数据与外部资源层", ["保存教务数据、AI配置数据、课程资料与任务结果，并连接外部模型服务。"]),
    ]
    for box, title, lines in layers:
        draw_box_text_centered(draw, box, title, lines, title_size=40, body_size=28)

    for x in (435, 1300, 2165):
        draw_arrow(draw, (x, 430), (x, 580))
    draw_arrow(draw, (1300, 790), (1300, 900))
    draw_arrow(draw, (1300, 1110), (1300, 1220))
    draw_arrow(draw, (1300, 1430), (1300, 1540))
    img.save(path)


def generate_closed_loop_figure(path: Path) -> None:
    img = Image.new("RGB", (2400, 780), "white")
    draw = ImageDraw.Draw(img)
    draw_center_text(draw, (1200, 70), "AI业务闭环逻辑", load_font(54, bold=True))

    boxes = [
        ((80, 220, 560, 520), "配置准备", ["管理员准备模型", "知识资料和场景参数"]),
        ((660, 220, 1140, 520), "场景调用", ["学生咨询", "课程答疑", "教案生成"]),
        ((1240, 220, 1720, 520), "结果沉淀", ["保存问答记录", "课程资料变化", "任务状态和生成结果"]),
        ((1820, 220, 2300, 520), "持续优化", ["补充知识内容", "调整模型与场景配置"]),
    ]
    centers = []
    for box, title, lines in boxes:
        draw_box_text_centered(draw, box, title, lines, title_size=42, body_size=28)
        centers.append(((box[0] + box[2]) // 2, (box[1] + box[3]) // 2))

    draw_arrow(draw, (560, centers[0][1]), (660, centers[1][1]))
    draw_arrow(draw, (1140, centers[1][1]), (1240, centers[2][1]))
    draw_arrow(draw, (1720, centers[2][1]), (1820, centers[3][1]))
    draw.line((2050, 520, 2050, 660, 320, 660, 320, 520), fill="black", width=5)
    draw_arrow(draw, (320, 660), (320, 520))
    draw_label(draw, (1185, 620), "运行数据回流到下一轮知识维护与策略调整", 24)
    img.save(path)


def generate_layered_figure(path: Path) -> None:
    img = Image.new("RGB", (2400, 1120), "white")
    draw = ImageDraw.Draw(img)
    draw_center_text(draw, (1200, 70), "AI分层设计逻辑", load_font(54, bold=True))

    boxes = [
        ((180, 170, 2220, 350), "第一层：AI配置编排层", ["围绕模型接入、场景工作流和参数策略进行统一配置，为不同AI场景提供可治理的底座。"]),
        ((180, 410, 2220, 590), "第二层：知识组织层", ["负责制度文档、课程资料和教学材料的整理、归类、更新和沉淀，保证知识来源清晰稳定。"]),
        ((180, 650, 2220, 830), "第三层：推理执行层", ["负责请求理解、知识匹配、上下文组织和结果生成，是问答与教案生成真正发生的处理层。"]),
        ((180, 890, 2220, 1070), "第四层：业务接入层", ["将统一能力开放给AI客服、AI课程助手和智能教案等页面入口，使不同角色能够按场景直接使用。"]),
    ]
    for box, title, lines in boxes:
        draw_box_text_centered(draw, box, title, lines, title_size=40, body_size=28)

    draw_arrow(draw, (1200, 350), (1200, 410))
    draw_arrow(draw, (1200, 590), (1200, 650))
    draw_arrow(draw, (1200, 830), (1200, 890))
    img.save(path)


def draw_entity_box(
    draw: ImageDraw.ImageDraw,
    box: tuple[int, int, int, int],
    title: str,
    fields: list[str],
) -> None:
    x1, y1, x2, y2 = box
    draw.rectangle(box, outline="black", fill="white", width=4)
    divider_y = y1 + 84
    draw.line((x1, divider_y, x2, divider_y), fill="black", width=4)

    title_font = load_font(34, bold=True)
    body_font = load_font(26)
    draw_center_text(draw, ((x1 + x2) // 2, y1 + 42), title, title_font)

    max_width = x2 - x1 - 26
    body_lines: list[str] = []
    for field in fields:
        body_lines.extend(wrap_text(draw, field, body_font, max_width))
    line_h = 34
    total_h = len(body_lines) * line_h
    y = divider_y + max(12, (y2 - divider_y - total_h) // 2) + line_h // 2 - 2
    for line in body_lines:
        draw_center_text(draw, ((x1 + x2) // 2, y), line, body_font)
        y += line_h


def generate_er_figure(path: Path) -> None:
    img = Image.new("RGB", (3600, 2500), "white")
    draw = ImageDraw.Draw(img)
    draw_center_text(draw, (1800, 80), "AI核心数据ER图", load_font(56, bold=True))

    boxes = {
        "sys_users": (120, 220, 860, 770),
        "courses": (120, 1460, 860, 2060),
        "ai_model_apis": (1000, 220, 1780, 880),
        "ai_knowledge_bases": (1000, 1280, 1780, 2080),
        "ai_workflow_apps": (1920, 220, 2700, 980),
        "ai_kb_documents": (1920, 1260, 2700, 1940),
        "ai_lesson_plan_tasks": (2840, 220, 3480, 980),
        "ai_kb_chunks": (2840, 1260, 3480, 2080),
    }

    draw_entity_box(draw, boxes["sys_users"], "sys_users", ["PK id", "username", "role", "is_active"])
    draw_entity_box(draw, boxes["courses"], "courses", ["PK id", "name", "teacher_id", "credit", "capacity", "course_type"])
    draw_entity_box(draw, boxes["ai_model_apis"], "ai_model_apis", ["PK id", "name", "provider", "model_name", "endpoint", "enabled", "is_default"])
    draw_entity_box(draw, boxes["ai_knowledge_bases"], "ai_knowledge_bases", ["PK id", "slug", "name", "owner_type", "owner_user_id", "course_id", "feature"])
    draw_entity_box(draw, boxes["ai_workflow_apps"], "ai_workflow_apps", ["PK id", "code", "type", "name", "knowledge_base_id", "model_api_id", "owner_user_id", "course_id", "status"])
    draw_entity_box(draw, boxes["ai_kb_documents"], "ai_kb_documents", ["PK id", "knowledge_base_id", "title", "original_filename", "url", "file_ext", "enabled"])
    draw_entity_box(draw, boxes["ai_lesson_plan_tasks"], "ai_lesson_plan_tasks", ["PK id", "teacher_user_id", "course_id", "title", "status", "result", "knowledge_base_id", "model_api_id"])
    draw_entity_box(draw, boxes["ai_kb_chunks"], "ai_kb_chunks", ["PK id", "knowledge_base_id", "document_id", "seq", "content", "tokens", "document_title"])

    relations = [
        ((860, 450), (1000, 450), "拥有/配置"),
        ((860, 1710), (1000, 1710), "课程关联"),
        ((1780, 520), (1920, 520), "模型绑定"),
        ((1780, 1440), (1920, 1440), "知识绑定"),
        ((2700, 1600), (2840, 1600), "1:N"),
        ((2700, 700), (2840, 700), "任务模型"),
    ]
    for start, end, label in relations:
        draw_arrow(draw, start, end)
        draw_label(draw, ((start[0] + end[0]) // 2, start[1] - 38), label, 24)

    draw.line((1390, 880, 1390, 1080, 2310, 1080, 2310, 1260), fill="black", width=5)
    draw_label(draw, (1850, 1030), "工作流使用知识库", 24)
    draw.line((2310, 1940, 2310, 2200, 3160, 2200, 3160, 2080), fill="black", width=5)
    draw_label(draw, (2735, 2150), "文档切分为知识分块", 24)
    draw.line((520, 770, 520, 1110, 3160, 1110, 3160, 980), fill="black", width=5)
    draw_label(draw, (1800, 1060), "教师创建教案任务", 24)

    img.save(path)


def paragraph_contains(doc: Document, needle: str) -> Paragraph:
    for para in doc.paragraphs:
        if needle in para.text:
            return para
    raise ValueError(f"Cannot find paragraph containing: {needle}")


def paragraph_index(doc: Document, target: Paragraph) -> int:
    for i, para in enumerate(doc.paragraphs):
        if para._element == target._element:
            return i
    raise ValueError("Paragraph index not found.")


def delete_paragraph(paragraph: Paragraph) -> None:
    element = paragraph._element
    parent = element.getparent()
    parent.remove(element)


def insert_before(anchor: Paragraph, text: str = "", style=None) -> Paragraph:
    para = anchor.insert_paragraph_before(text)
    if style is not None:
        para.style = style
    return para


def replace_section_paragraphs(doc: Document, heading_text: str, stop_anchor_text: str, texts: list[str]) -> None:
    heading = paragraph_contains(doc, heading_text)
    anchor = paragraph_contains(doc, stop_anchor_text)
    start_idx = paragraph_index(doc, heading)
    end_idx = paragraph_index(doc, anchor)
    old_paras = doc.paragraphs[start_idx + 1:end_idx]
    for para in old_paras:
        delete_paragraph(para)

    normal_style = doc.styles["Normal"]
    for text in texts:
        insert_before(anchor, text, normal_style)


def insert_usecase_before_table(doc: Document, image_path: Path) -> None:
    table_caption = paragraph_contains(doc, "角色需求与AI价值对应关系")
    image_style = paragraph_contains(doc, "系统总体架构图")._element.getprevious()
    # pull style/alignment from an existing image paragraph
    existing_image_para = doc.paragraphs[paragraph_index(doc, paragraph_contains(doc, "系统总体架构图")) - 1]
    image_para = insert_before(table_caption, style=existing_image_para.style)
    image_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    image_para.add_run().add_picture(str(image_path), width=Cm(16))

    caption_para = insert_before(table_caption, "图二-1 系统角色需求用例图", paragraph_contains(doc, "系统总体架构图").style)
    caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER


def replace_picture_before_caption(doc: Document, caption_keyword: str, image_path: Path, width_cm: float = 16.0) -> None:
    caption_para = paragraph_contains(doc, caption_keyword)
    pic_para = doc.paragraphs[paragraph_index(doc, caption_para) - 1]
    pic_para.clear()
    pic_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pic_para.add_run().add_picture(str(image_path), width=Cm(width_cm))


def set_paragraph_text_by_keyword(doc: Document, keyword: str, text: str) -> None:
    para = paragraph_contains(doc, keyword)
    para.clear()
    para.add_run(text)
    if "图" in text:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER


def update_chapter_two_text(doc: Document) -> None:
    replace_section_paragraphs(
        doc,
        "业务场景与角色需求分析",
        "角色需求与AI价值对应关系",
        [
            "第二章的需求分析不是先假设系统应该长成什么样，再反过来给它安上几个AI功能，而是从现有教务系统的实际使用过程往回推。学生、教师和管理员虽然都在同一平台中活动，但他们面对的问题性质并不一样：学生更多是查、问、办，教师更多是管、教、备，管理员更多是配、控、维。这种差异，决定了需求分析不能只罗列角色，更要解释为什么这些角色会形成当前这样的需求结构。",
            "先看学生侧。传统教务系统在信息组织上通常比较规范，菜单、表单和制度说明都不缺，但真正使用时，学生遇到的麻烦往往不是“没有入口”，而是“看见入口以后还是不明白怎么走下一步”。例如选课信息能查到，请假页面也能打开，办事大厅也有事项列表，可一旦涉及办理条件、材料差异、流程顺序或者课程内容理解，静态页面就很难继续承担解释工作。这正是学生侧需要引入自然语言帮助的直接原因。",
            "进一步说，学生在系统里的操作通常是连续的。查看课程信息只是开始，后面可能还要决定是否办理、去哪个页面继续提交、遇到疑问该向谁确认；学习场景也是一样，看到课程入口后，真正的问题往往出现在知识理解阶段。也就是说，学生需求不是孤立地需要一个“聊天框”，而是希望在原有业务链条上多一个能及时解释、能降低理解成本、又不会打断办理过程的辅助层。",
            "教师侧的需求来源则更复杂一些。教师在教务系统中的工作，一部分是确定性的，如课程管理、成绩处理、审批等；另一部分却高度依赖经验和内容组织，比如整理课程资料、反复回答共性问题、准备教案和教学提纲。前一部分传统教务系统已经覆盖得比较成熟，后一部分长期依赖人工完成，工作量零散却持续。正因为这种结构性差异，教师对AI的需求天然会落在“围绕课程内容减轻重复劳动”上。",
            "从实际教学过程看，教师不会接受一个完全脱离课程边界的通用问答工具。原因很简单：课程答疑要对应教材、课件、授课进度和本人的讲授重点，教案生成也必须能回到本课程资料上继续修改。因此，教师侧之所以会形成课程助手、资料沉淀和智能教案这些需求，并不是因为功能看起来先进，而是因为教学责任要求结果必须可追溯、可编辑、可继续复用。",
            "管理员侧的需求又是另一条线。只要系统把AI能力真正放进教务平台，就会立刻出现治理问题：模型能不能连通、知识资料归谁维护、哪些场景对哪些角色开放、配置改动后前台是否同步生效。如果这些问题不先解决，前面的学生咨询和教师备课支持就很容易变成一次性的演示功能，而不是可持续运行的系统能力。因此，管理员的需求本质上来自平台治理，而不是直接来自业务使用。",
            "把三类角色放在一起看，可以更清楚地看到需求是如何被推出来的。学生侧不断产生咨询与理解需求，教师侧不断沉淀课程资料和教学内容，管理员侧则负责保证底层AI资源始终可配、可控、可维护。三者之间既有边界，也有依赖：没有教师资料，课程问答会变空；没有管理员治理，模型和知识库难以稳定服务；没有学生和教师的实际使用，AI能力又缺少真实业务场景。这种相互依赖关系，正是本系统需求分析的核心依据。",
            "基于以上分析，系统最终可以归纳出三条比较稳定的需求主线：第一，面向学生，形成“查询之后还能继续理解和办理”的辅助能力；第二，面向教师，形成“围绕课程沉淀资料并持续复用”的教学支持能力；第三，面向管理员，形成“能够统一配置并稳定治理”的AI底座能力。图二-1将这种角色与需求之间的对应关系做了更直观的表达。",
        ],
    )

    replace_section_paragraphs(
        doc,
        "功能需求分析",
        "非功能需求分析",
        [
            "功能需求并不是从技术模块出发拆分出来的，而是从前一节已经明确的角色压力和业务链条中一步步抽象出来的。换句话说，本系统为什么需要这些功能，不是因为AI能够做这些事，而是因为学生、教师和管理员在现有教务流程中已经暴露出对应的问题，系统需要用较小的结构改动去把这些问题接住。",
            "首先，基础教务功能必须保留为系统的主骨架。这一点并不是保守设计，而是需求本身决定的。学生的选课、请假、成绩查询、证书查看和办事事项办理，教师的课程管理、作业管理和成绩处理，管理员的用户维护和流程配置，构成了整个系统的真实业务边界。AI只有依附在这些边界之内，回答和生成的内容才有明确场景，不会脱离教务流程空转。",
            "其次，AI客服的功能需求来自高频但分散的咨询场景。学生和教师遇到的问题，很多并不值得专门找人处理，却又确实会影响后续操作，例如事项该找哪个入口、材料是否齐全、某项流程先后顺序是什么、规则解释如何理解。由此可以推导出，AI客服不能只是“能回答”，还需要具备推荐问题、状态反馈、知识不足提示和业务入口引导等功能，才能真正服务于咨询场景，而不是停留在泛化问答层面。",
            "AI课程助手的功能需求则是从教学答疑链中推出来的。教师在课程教学中经常面对重复问题，但这些问题又不能完全交给通用知识回答，因为不同课程、不同资料、不同教学重点之间差异很大。因此，系统需要支持教师围绕课程上传资料、形成课程专属知识来源，并让学生在明确的课程入口下发起提问。课程绑定、资料更新、助手选择和结果反馈这些功能，看起来是页面操作，实际上都是从“答疑必须贴合具体课程”这一需求前提中自然派生出来的。",
            "智能教案的功能需求来源于备课场景的重复性和迭代性。教师真正需要的不是一次性吐出一段文本，而是能围绕课程主题、资料基础和教学目标形成一份初稿，并在后续继续修改、保存和导出。正因为备课本身是一个多轮调整过程，系统才需要把教案生成设计成带有任务状态、结果回写、历史保留和再次编辑能力的功能，而不是把它处理成一次性的即时回答。",
            "管理端AI治理功能同样不是附属功能，而是由平台运行需求倒推出来的。模型接口管理、知识库管理、工作流编排、客服参数维护以及连通性校验，分别对应“能不能接入”“资料由谁维护”“不同场景如何分流”“最终展示什么内容”“配置改完能不能立即验证”这些很实际的问题。少了这一层，前台的AI入口很难长期稳定使用。",
            "除此之外，系统还需要一组跨模块的支撑功能。这些功能表面上不如AI客服、课程助手和智能教案显眼，但实际上决定了整个链条是否能闭合，包括角色权限校验、上传与生成状态提示、历史记录保存、结果再次复用以及异常信息反馈等。它们并不是额外附加的“优化项”，而是从完整业务链条中倒推出来的必需能力。也正因为有了这些支撑功能，系统中的AI模块才不是孤立页面，而是能与现有教务业务协同运行的整体能力。",
        ],
    )


def update_figure_captions(doc: Document) -> None:
    set_paragraph_text_by_keyword(doc, "系统总体架构图", "图二-2 系统总体架构图")
    set_paragraph_text_by_keyword(doc, "AI业务闭环逻辑", "图二-3 AI业务闭环逻辑")
    set_paragraph_text_by_keyword(doc, "AI分层设计逻辑", "图二-4 AI分层设计逻辑")
    set_paragraph_text_by_keyword(doc, "AI核心数据ER图", "图三-1 AI核心数据ER图")


def main() -> None:
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    input_doc = get_input_doc()
    output_doc = get_output_doc(input_doc)
    ensure_backup(input_doc)

    fig_usecase = ASSET_DIR / "fig_2_1_usecase_uml_black.png"
    fig_arch = ASSET_DIR / "fig_2_2_system_architecture_black.png"
    fig_loop = ASSET_DIR / "fig_2_3_ai_closed_loop_black.png"
    fig_layer = ASSET_DIR / "fig_2_4_ai_layered_black.png"
    fig_er = ASSET_DIR / "fig_3_1_ai_er_black.png"

    generate_usecase_figure(fig_usecase)
    generate_system_architecture_figure(fig_arch)
    generate_closed_loop_figure(fig_loop)
    generate_layered_figure(fig_layer)
    generate_er_figure(fig_er)

    doc = Document(input_doc)
    update_chapter_two_text(doc)
    insert_usecase_before_table(doc, fig_usecase)

    replace_picture_before_caption(doc, "系统总体架构图", fig_arch)
    replace_picture_before_caption(doc, "AI业务闭环逻辑", fig_loop)
    replace_picture_before_caption(doc, "AI分层设计逻辑", fig_layer)
    replace_picture_before_caption(doc, "AI核心数据ER图", fig_er)

    update_figure_captions(doc)

    if output_doc.exists():
        output_doc.unlink()
    doc.save(output_doc)
    print(output_doc)


if __name__ == "__main__":
    main()
