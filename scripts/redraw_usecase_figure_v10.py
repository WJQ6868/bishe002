from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import Cm
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
ASSET_DIR = ROOT / "generated" / "thesis_assets"
FONT_REGULAR = Path(r"C:\Windows\Fonts\msyh.ttc")
FONT_BOLD = Path(r"C:\Windows\Fonts\simhei.ttf")


def load_font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont:
    font_path = FONT_BOLD if bold and FONT_BOLD.exists() else FONT_REGULAR
    return ImageFont.truetype(str(font_path), size=size)


def wrap_text(draw: ImageDraw.ImageDraw, text: str, font: ImageFont.FreeTypeFont, max_width: int) -> list[str]:
    lines: list[str] = []
    current = ""
    for ch in text:
        candidate = current + ch
        if draw.textlength(candidate, font=font) <= max_width:
            current = candidate
        else:
            if current:
                lines.append(current)
            current = ch
    if current:
        lines.append(current)
    return lines or [text]


def draw_center_text(draw: ImageDraw.ImageDraw, center: tuple[int, int], text: str, font, fill="black") -> None:
    bbox = draw.textbbox((0, 0), text, font=font)
    w = bbox[2] - bbox[0]
    h = bbox[3] - bbox[1]
    draw.text((center[0] - w / 2, center[1] - h / 2), text, font=font, fill=fill)


def draw_actor(draw: ImageDraw.ImageDraw, x: int, y: int, label: str) -> None:
    draw.ellipse((x - 18, y, x + 18, y + 36), outline="black", width=4)
    draw.line((x, y + 36, x, y + 112), fill="black", width=4)
    draw.line((x - 38, y + 68, x + 38, y + 68), fill="black", width=4)
    draw.line((x, y + 112, x - 32, y + 162), fill="black", width=4)
    draw.line((x, y + 112, x + 32, y + 162), fill="black", width=4)
    draw_center_text(draw, (x, y + 214), label, load_font(34, bold=True))


def draw_usecase(draw: ImageDraw.ImageDraw, box: tuple[int, int, int, int], text: str) -> None:
    x1, y1, x2, y2 = box
    draw.ellipse(box, outline="black", fill="white", width=4)
    font = load_font(32, bold=False)
    lines = wrap_text(draw, text, font, x2 - x1 - 60)
    line_h = 42
    total_h = len(lines) * line_h
    y = y1 + (y2 - y1 - total_h) // 2 + line_h // 2 - 2
    cx = (x1 + x2) // 2
    for line in lines:
        draw_center_text(draw, (cx, y), line, font)
        y += line_h


def draw_line(draw: ImageDraw.ImageDraw, points: list[tuple[int, int]], width: int = 4) -> None:
    if len(points) >= 2:
        draw.line(points, fill="black", width=width)


def generate_figure(path: Path) -> None:
    img = Image.new("RGB", (3000, 1800), "white")
    draw = ImageDraw.Draw(img)

    draw_center_text(draw, (1500, 70), "系统角色需求用例图", load_font(54, bold=True))

    boundary = (650, 150, 2350, 1520)
    draw.rectangle(boundary, outline="black", width=5)
    draw.text((680, 180), "AI赋能高校教务系统", font=load_font(34, bold=True), fill="black")

    draw_actor(draw, 360, 340, "学生")
    draw_actor(draw, 360, 860, "教师")
    draw_actor(draw, 2640, 600, "管理员")

    usecases = {
        "course_manage": ((880, 360, 1380, 540), "课程查询与办理"),
        "ai_service": ((1660, 220, 2160, 400), "AI客服咨询"),
        "course_qa": ((1260, 600, 1780, 790), "课程答疑支持"),
        "upload": ((880, 900, 1480, 1080), "课程资料上传与沉淀"),
        "lesson": ((1600, 1060, 2200, 1240), "智能教案生成与修改"),
        "workflow": ((880, 1290, 1480, 1490), "工作流编排与治理"),
        "model": ((1600, 1290, 2200, 1490), "模型与知识库配置"),
    }

    for box, text in usecases.values():
        draw_usecase(draw, box, text)

    # Student associations
    draw_line(draw, [(398, 405), (880, 450)])
    draw_line(draw, [(398, 380), (1660, 310)])
    draw_line(draw, [(398, 430), (1260, 690)])

    # Teacher associations
    draw_line(draw, [(398, 900), (1260, 690)])
    draw_line(draw, [(398, 925), (880, 990)])
    draw_line(draw, [(398, 1020), (1600, 1150)])

    # Administrator associations
    draw_line(draw, [(2602, 640), (2200, 1150)])
    draw_line(draw, [(2602, 665), (1480, 1390)])
    draw_line(draw, [(2602, 690), (2200, 1390)])

    img.save(path)


def replace_image_in_doc(doc_path: Path, out_path: Path, image_path: Path) -> None:
    doc = Document(doc_path)
    caption_idx = next(i for i, p in enumerate(doc.paragraphs) if "图二-1 系统角色需求用例图" in p.text)
    image_para = doc.paragraphs[caption_idx - 1]
    image_para.clear()
    image_para.add_run().add_picture(str(image_path), width=Cm(16))
    if out_path.exists():
        out_path.unlink()
    doc.save(out_path)


def main() -> None:
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    desktop = Path.home() / "Desktop"
    doc_path = next(desktop.glob("2405273202*chapter2-updated-v9.docx"))
    out_path = doc_path.with_name(doc_path.stem.replace("-v9", "-v12") + doc_path.suffix)
    image_path = ASSET_DIR / "fig_2_1_usecase_uml_black_v2.png"
    generate_figure(image_path)
    replace_image_in_doc(doc_path, out_path, image_path)
    print(out_path)
    print(image_path)


if __name__ == "__main__":
    main()
