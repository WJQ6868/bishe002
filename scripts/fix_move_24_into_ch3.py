from pathlib import Path
import shutil

from docx import Document


DESKTOP_DOC = Path(
    r"C:\Users\wangj\Desktop\2405273202-王佳齐-AI赋能的高校教务系统ai.docx"
)
WORK_DOC = Path(r"D:\bishe\one\paper_ch24_moved_to_ch3.docx")
BACKUP_DOC = Path(
    r"C:\Users\wangj\Desktop\2405273202-王佳齐-AI赋能的高校教务系统ai-before-fix-2_4-move-20260520.docx"
)


def set_paragraph_text(paragraph, text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def replace_if_startswith(doc: Document, startswith: str, new_text: str) -> int:
    count = 0
    for paragraph in doc.paragraphs:
        if paragraph.text.startswith(startswith):
            set_paragraph_text(paragraph, new_text)
            count += 1
    return count


def main() -> None:
    if not WORK_DOC.exists():
        raise FileNotFoundError(f"Work doc not found: {WORK_DOC}")

    doc = Document(str(WORK_DOC))

    expected_prefixes = {
        140: ["\u5168\u6587\u5171\u5206\u4e3a\u516d\u7ae0\u3002"],
        186: ["??????????", "\u975e\u529f\u80fd\u4e0e\u8fd0\u884c\u4fdd\u969c\u8bbe\u8ba1"],
        187: ["?????", "\u53ef\u7528\u6027\u8bbe\u8ba1"],
        189: ["??????", "\u8fde\u7eed\u53cd\u9988\u8bbe\u8ba1"],
        191: ["????????", "\u53ef\u9760\u6027\u4e0e\u8ffd\u8e2a\u8bbe\u8ba1"],
        193: ["?????????", "\u5b89\u5168\u4e0e\u6743\u9650\u63a7\u5236\u8bbe\u8ba1"],
        195: ["???????", "\u90e8\u7f72\u4e0e\u8fd0\u7ef4\u8bbe\u8ba1"],
        197: ["AI????????", "AI\u80fd\u529b\u63a5\u5165\u8fb9\u754c\u8bbe\u8ba1"],
        198: ["??????", "\u63a5\u5165\u8fb9\u754c\u8bbe\u8ba1"],
        200: ["???AI????", "\u4e1a\u52a1\u4e0eAI\u534f\u540c\u8bbe\u8ba1"],
        203: ["?3-2", "\u56fe3-2"],
        215: ["?3-3", "\u56fe3-3"],
        219: ["?3-4", "\u56fe3-4"],
    }
    for idx, prefixes in expected_prefixes.items():
        actual = doc.paragraphs[idx].text
        if not any(actual.startswith(prefix) for prefix in prefixes):
            raise ValueError(f"Paragraph {idx} changed unexpectedly: {actual!r}")

    set_paragraph_text(
        doc.paragraphs[140],
        "\u5168\u6587\u5171\u5206\u4e3a\u516d\u7ae0\u3002\u7b2c\u4e00\u7ae0\u4e3b\u8981\u4ecb\u7ecd\u8bfe\u9898\u80cc\u666f\u3001\u7814\u7a76\u610f\u4e49\u3001\u56fd\u5185\u5916\u7814\u7a76\u73b0\u72b6\u4ee5\u53ca\u672c\u6587\u7684\u7814\u7a76\u76ee\u6807\u4e0e\u6280\u672f\u8def\u7ebf\uff0c\u4e3a\u540e\u7eed\u7814\u7a76\u5185\u5bb9\u5960\u5b9a\u57fa\u7840\u3002\u7b2c\u4e8c\u7ae0\u56f4\u7ed5\u9ad8\u6821\u6559\u52a1\u573a\u666f\u5c55\u5f00\u9700\u6c42\u5206\u6790\uff0c\u91cd\u70b9\u8bf4\u660e\u5b66\u751f\u3001\u6559\u5e08\u548c\u7ba1\u7406\u5458\u4e09\u7c7b\u89d2\u8272\u7684\u4f7f\u7528\u9700\u6c42\uff0c\u5e76\u4ece\u4e1a\u52a1\u573a\u666f\u4e0e\u529f\u80fd\u9700\u6c42\u4e24\u4e2a\u5c42\u9762\u8fdb\u884c\u68b3\u7406\u3002\u7b2c\u4e09\u7ae0\u4e3a\u7cfb\u7edf\u8bbe\u8ba1\u7ae0\u8282\uff0c\u4e3b\u8981\u5bf9\u7cfb\u7edf\u6574\u4f53\u529f\u80fd\u7ec4\u7ec7\u3001\u975e\u529f\u80fd\u4e0e\u8fd0\u884c\u4fdd\u969c\u3001AI\u80fd\u529b\u63a5\u5165\u8fb9\u754c\u3001\u7cfb\u7edf\u534f\u540c\u5173\u7cfb\u3001\u6570\u636e\u7ec4\u7ec7\u3001AI\u95ed\u73af\u8fd0\u884c\u3001\u6280\u672f\u652f\u6491\u3001\u5de5\u4f5c\u6d41\u4e0e\u77e5\u8bc6\u8d44\u6e90\u7ec4\u7ec7\u3001\u8f7b\u91cf\u7ea7\u68c0\u7d22\u589e\u5f3a\u751f\u6210\u65b9\u6848\u3001\u6a21\u578b\u9009\u62e9\u7b56\u7565\u3001SSE\u6d41\u5f0f\u54cd\u5e94\u673a\u5236\u4ee5\u53ca\u6838\u5fc3\u6570\u636e\u8868\u7ed3\u6784\u8fdb\u884c\u8bbe\u8ba1\u8bf4\u660e\u3002\u7b2c\u56db\u7ae0\u4e3a\u7cfb\u7edf\u5b9e\u73b0\u7ae0\u8282\uff0c\u7ed3\u5408\u9879\u76ee\u5b9e\u9645\u4ee3\u7801\uff0c\u5bf9\u7528\u6237\u3001\u8bfe\u7a0b\u3001\u9009\u8bfe\u3001\u8bf7\u5047\u3001\u529e\u4e8b\u3001\u4f5c\u4e1a\u7b49\u57fa\u7840\u6559\u52a1\u529f\u80fd\u7684\u5b9e\u73b0\u65b9\u5f0f\u4f5c\u7b80\u8981\u8bf4\u660e\uff0c\u5e76\u91cd\u70b9\u9610\u8ff0AI\u5ba2\u670d\u3001AI\u8bfe\u7a0b\u52a9\u624b\u3001\u667a\u80fd\u6559\u6848\u548c\u7ba1\u7406\u7aefAI\u914d\u7f6e\u7b49\u6a21\u5757\u7684\u9875\u9762\u4ea4\u4e92\u3001\u63a5\u53e3\u8c03\u7528\u3001\u540e\u7aef\u5904\u7406\u548c\u6570\u636e\u843d\u5e93\u8fc7\u7a0b\u3002\u7b2c\u4e94\u7ae0\u5bf9\u7cfb\u7edf\u6d4b\u8bd5\u73af\u5883\u3001\u6d4b\u8bd5\u7528\u4f8b\u3001\u6d4b\u8bd5\u7ed3\u679c\u548c\u529f\u80fd\u8fd0\u884c\u6548\u679c\u8fdb\u884c\u5206\u6790\uff0c\u9a8c\u8bc1\u7cfb\u7edf\u4e3b\u8981AI\u529f\u80fd\u94fe\u8def\u7684\u53ef\u7528\u6027\u3002\u7b2c\u516d\u7ae0\u5bf9\u5168\u6587\u7814\u7a76\u5185\u5bb9\u8fdb\u884c\u603b\u7ed3\uff0c\u5e76\u5bf9\u7cfb\u7edf\u540e\u7eed\u4f18\u5316\u65b9\u5411\u8fdb\u884c\u5c55\u671b\u3002",
    )

    replacements = {
        186: "\u975e\u529f\u80fd\u4e0e\u8fd0\u884c\u4fdd\u969c\u8bbe\u8ba1",
        187: "\u53ef\u7528\u6027\u8bbe\u8ba1",
        189: "\u8fde\u7eed\u53cd\u9988\u8bbe\u8ba1",
        191: "\u53ef\u9760\u6027\u4e0e\u8ffd\u8e2a\u8bbe\u8ba1",
        193: "\u5b89\u5168\u4e0e\u6743\u9650\u63a7\u5236\u8bbe\u8ba1",
        195: "\u90e8\u7f72\u4e0e\u8fd0\u7ef4\u8bbe\u8ba1",
        197: "AI\u80fd\u529b\u63a5\u5165\u8fb9\u754c\u8bbe\u8ba1",
        198: "\u63a5\u5165\u8fb9\u754c\u8bbe\u8ba1",
        200: "\u4e1a\u52a1\u4e0eAI\u534f\u540c\u8bbe\u8ba1",
        203: "\u56fe3-2 \u57fa\u4e8e\u73b0\u6709\u6559\u52a1\u7cfb\u7edf\u7684AI\u589e\u5f3a\u603b\u4f53\u67b6\u6784\u56fe",
        215: "\u56fe3-3 AI\u4e1a\u52a1\u95ed\u73af\u903b\u8f91",
        219: "\u56fe3-4 AI\u8fd0\u884c\u652f\u6491\u5173\u7cfb\u56fe",
    }
    for idx, text in replacements.items():
        set_paragraph_text(doc.paragraphs[idx], text)

    for paragraph in doc.paragraphs:
        if paragraph.text.startswith(
            "\u5982\u679c\u76f4\u63a5\u4f9d\u8d56\u901a\u7528AI\u5de5\u5177"
        ):
            paragraph.style = doc.styles["Normal"]
        if "\u88682-2\u6240\u793a" in paragraph.text:
            set_paragraph_text(
                paragraph,
                paragraph.text.replace("\u88682-2\u6240\u793a", "\u5982\u4e0b\u8868\u6240\u793a"),
            )
        if "\u5982\u5982\u4e0b\u8868\u6240\u793a" in paragraph.text:
            set_paragraph_text(
                paragraph,
                paragraph.text.replace("\u5982\u5982\u4e0b\u8868\u6240\u793a", "\u5982\u4e0b\u8868\u6240\u793a"),
            )

    doc.save(str(WORK_DOC))

    if DESKTOP_DOC.exists():
        shutil.copy2(DESKTOP_DOC, BACKUP_DOC)
    shutil.copy2(WORK_DOC, DESKTOP_DOC)

    print(f"Updated: {WORK_DOC}")
    print(f"Backup: {BACKUP_DOC}")
    print(f"Copied to: {DESKTOP_DOC}")


if __name__ == "__main__":
    main()
