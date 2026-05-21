from __future__ import annotations

from pathlib import Path

from docx import Document


def main() -> None:
    desktop = Path.home() / "Desktop"
    input_doc = next(desktop.glob("2405273202*chapter2-updated-v6.docx"))
    output_doc = input_doc.with_name(input_doc.stem.replace("-v6", "-v7") + input_doc.suffix)
    doc = Document(input_doc)

    arch_caption_idx = next(i for i, p in enumerate(doc.paragraphs) if "系统总体架构图" in p.text)
    image_style = doc.paragraphs[arch_caption_idx - 1].style
    caption_style = doc.paragraphs[arch_caption_idx].style
    table_style = next(p.style for p in doc.paragraphs if "表二-2 系统模块划分与职责边界" in p.text or "表二-2系统模块划分与职责边界" in p.text)

    pic_para = doc.paragraphs[95]
    fig_cap = doc.paragraphs[96]
    table_cap = doc.paragraphs[97]

    pic_para.style = image_style
    fig_cap.clear()
    fig_cap.add_run("图二-1 系统角色需求用例图")
    fig_cap.style = caption_style

    table_cap.clear()
    table_cap.add_run("表二-1角色需求与AI价值对应关系")
    table_cap.style = table_style

    if output_doc.exists():
        output_doc.unlink()
    doc.save(output_doc)
    print(output_doc)


if __name__ == "__main__":
    main()
