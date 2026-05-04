from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parents[1]
SOURCE_DOCX = ROOT / "AI赋能的高校教务系统-毕业论文润色定稿版.docx"
OUTPUT_DOCX = ROOT / "AI赋能的高校教务系统-毕业论文模板规范版.docx"


RE_CHAPTER = re.compile(r"^第[一二三四五六七八九十]+章\s*")
RE_SEC2 = re.compile(r"^\d+\.\d+\s+")
RE_SEC3 = re.compile(r"^\d+\.\d+\.\d+\s+")
RE_FIG = re.compile(r"^图\d+[-－]\d+")
RE_TAB = re.compile(r"^表\d+[-－]\d+")
RE_REF = re.compile(r"^\[\d+\]")


def set_run_font(run, *, east_asia: str = "宋体", ascii_font: str = "Times New Roman", size_pt: float = 12, bold: bool | None = None):
    run.font.name = ascii_font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    run.font.size = Pt(size_pt)
    if bold is not None:
        run.font.bold = bold


def clear_paragraph_format(p):
    pf = p.paragraph_format
    pf.left_indent = None
    pf.right_indent = None
    pf.first_line_indent = None
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = 1.5
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE


def format_cover_center(p, size_pt: float, *, east_asia: str = "黑体", bold: bool = True):
    clear_paragraph_format(p)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        set_run_font(run, east_asia=east_asia, size_pt=size_pt, bold=bold)


def format_title_paragraph(p, text: str):
    clear_paragraph_format(p)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    for run in p.runs:
        if run.text.strip():
            if text == "ABSTRACT":
                set_run_font(run, east_asia="Times New Roman", ascii_font="Times New Roman", size_pt=15, bold=True)
            else:
                set_run_font(run, east_asia="黑体", size_pt=15, bold=True)


def format_body_paragraph(p):
    clear_paragraph_format(p)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Pt(24)
    for run in p.runs:
        set_run_font(run, east_asia="宋体", size_pt=12, bold=False)


def format_heading1(p):
    clear_paragraph_format(p)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)
    for run in p.runs:
        set_run_font(run, east_asia="黑体", size_pt=15, bold=True)


def format_heading2(p):
    clear_paragraph_format(p)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        set_run_font(run, east_asia="黑体", size_pt=14, bold=True)


def format_heading3(p):
    clear_paragraph_format(p)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        set_run_font(run, east_asia="黑体", size_pt=12, bold=True)


def format_keywords_cn(p):
    clear_paragraph_format(p)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.first_line_indent = Pt(0)
    if not p.runs:
        return
    for run in p.runs:
        if "关键词" in run.text:
            set_run_font(run, east_asia="黑体", size_pt=15, bold=False)
        else:
            set_run_font(run, east_asia="宋体", size_pt=12, bold=False)


def format_keywords_en(p):
    clear_paragraph_format(p)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.first_line_indent = Pt(0)
    for run in p.runs:
        if "Key words" in run.text:
            set_run_font(run, east_asia="Times New Roman", ascii_font="Times New Roman", size_pt=15, bold=False)
        else:
            set_run_font(run, east_asia="Times New Roman", ascii_font="Times New Roman", size_pt=12, bold=False)


def format_caption(p, *, is_table: bool):
    clear_paragraph_format(p)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    if is_table:
        p.paragraph_format.space_before = Pt(5)
    else:
        p.paragraph_format.space_after = Pt(5)
    for run in p.runs:
        set_run_font(run, east_asia="宋体", size_pt=10.5, bold=False)


def format_reference(p):
    clear_paragraph_format(p)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.first_line_indent = Pt(-21)
    p.paragraph_format.left_indent = Pt(21)
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    for run in p.runs:
        set_run_font(run, east_asia="宋体", size_pt=10.5, bold=False)


def set_cell_text_style(cell, *, east_asia: str = "宋体", size_pt: float = 10.5, bold_first_row: bool = False, row_idx: int = 0):
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        clear_paragraph_format(paragraph)
        paragraph.paragraph_format.first_line_indent = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.0
        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        for run in paragraph.runs:
            set_run_font(run, east_asia=east_asia, size_pt=size_pt, bold=(bold_first_row and row_idx == 0))


def set_table_borders(table, *, cover_table: bool = False):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)

    def make_edge(name: str, val: str = "nil", sz: str = "4"):
        edge = borders.find(qn(f"w:{name}"))
        if edge is None:
            edge = OxmlElement(f"w:{name}")
            borders.append(edge)
        edge.set(qn("w:val"), val)
        if val != "nil":
            edge.set(qn("w:sz"), sz)
            edge.set(qn("w:space"), "0")
            edge.set(qn("w:color"), "000000")

    if cover_table:
        for name in ("top", "left", "bottom", "right", "insideH", "insideV"):
            make_edge(name, "nil")
        return

    make_edge("top", "single", "8")
    make_edge("bottom", "single", "8")
    make_edge("left", "nil")
    make_edge("right", "nil")
    make_edge("insideV", "nil")
    make_edge("insideH", "nil")

    if table.rows:
        for cell in table.rows[0].cells:
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_borders = tc_pr.first_child_found_in("w:tcBorders")
            if tc_borders is None:
                tc_borders = OxmlElement("w:tcBorders")
                tc_pr.append(tc_borders)
            bottom = tc_borders.find(qn("w:bottom"))
            if bottom is None:
                bottom = OxmlElement("w:bottom")
                tc_borders.append(bottom)
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "8")
            bottom.set(qn("w:space"), "0")
            bottom.set(qn("w:color"), "000000")


def format_table(table, *, cover_table: bool = False):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table, cover_table=cover_table)
    for r_idx, row in enumerate(table.rows):
        for cell in row.cells:
            set_cell_text_style(cell, east_asia="宋体", size_pt=(12 if cover_table else 10.5), bold_first_row=not cover_table, row_idx=r_idx)


def set_section_layout(doc: Document):
    for sec in doc.sections:
        sec.page_width = Cm(21.0)
        sec.page_height = Cm(29.7)
        sec.top_margin = Cm(2.5)
        sec.bottom_margin = Cm(2.0)
        sec.left_margin = Cm(2.0)
        sec.right_margin = Cm(2.0)
        sec.header_distance = Cm(2.0)
        sec.footer_distance = Cm(1.5)
        for p in sec.footer.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                set_run_font(run, east_asia="Times New Roman", ascii_font="Times New Roman", size_pt=10.5, bold=False)


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--source", type=Path, default=SOURCE_DOCX)
    parser.add_argument("--output", type=Path, default=OUTPUT_DOCX)
    args = parser.parse_args()

    source = args.source
    output = args.output

    doc = Document(str(source))
    set_section_layout(doc)

    in_references = False
    for p in doc.paragraphs:
        text = (p.text or "").strip()
        if not text:
            continue

        if text in {"南京工业职业技术大学"}:
            format_cover_center(p, 22)
            continue
        if text in {"本科毕业设计（论文）"}:
            format_cover_center(p, 22)
            continue
        if text == "AI赋能的高校教务系统":
            format_cover_center(p, 22)
            continue
        if re.fullmatch(r"\d{4}年\d{1,2}月", text):
            format_cover_center(p, 15, east_asia="黑体", bold=False)
            continue

        if text in {"诚信承诺书", "摘  要", "目  录", "参考文献", "致谢"}:
            format_title_paragraph(p, text)
            in_references = text == "参考文献"
            continue
        if text == "ABSTRACT":
            format_title_paragraph(p, text)
            continue

        if RE_CHAPTER.match(text):
            in_references = False
            format_heading1(p)
            continue
        if RE_SEC3.match(text):
            in_references = False
            format_heading3(p)
            continue
        if RE_SEC2.match(text):
            in_references = False
            format_heading2(p)
            continue
        if RE_FIG.match(text):
            format_caption(p, is_table=False)
            continue
        if RE_TAB.match(text):
            format_caption(p, is_table=True)
            continue
        if text.startswith("关键词："):
            format_keywords_cn(p)
            continue
        if text.startswith("Key words:"):
            format_keywords_en(p)
            continue
        if in_references and RE_REF.match(text):
            format_reference(p)
            continue

        if text in {"附录", "附录A", "附录 B"} or text.startswith("附录"):
            format_heading1(p)
            continue

        format_body_paragraph(p)

    for idx, table in enumerate(doc.tables):
        format_table(table, cover_table=(idx == 0))

    doc.save(str(output))
    print(output)


if __name__ == "__main__":
    main()
