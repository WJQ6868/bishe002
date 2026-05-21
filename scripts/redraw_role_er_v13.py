from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import Cm
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
ASSET_DIR = ROOT / "generated" / "thesis_assets"
FONT_REGULAR = Path(r"C:\Windows\Fonts\msyh.ttc")
FONT_BOLD = Path(r"C:\Windows\Fonts\simhei.ttf")


def load_font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont:
    font_path = FONT_BOLD if bold and FONT_BOLD.exists() else FONT_REGULAR
    return ImageFont.truetype(str(font_path), size=size)


def wrap_text(draw: ImageDraw.ImageDraw, text: str, font: ImageFont.FreeTypeFont, max_width: int) -> list[str]:
    if not text:
        return [""]
    lines: list[str] = []
    current = ""
    for ch in text:
        candidate = current + ch
        if draw.textlength(candidate, font=font) <= max_width:
            current = candidate
        else:
            if current:
                lines.append(current)
            current = ch
    if current:
        lines.append(current)
    return lines or [text]


def draw_center_text(draw: ImageDraw.ImageDraw, center: tuple[int, int], text: str, font, fill="black") -> None:
    bbox = draw.textbbox((0, 0), text, font=font)
    w = bbox[2] - bbox[0]
    h = bbox[3] - bbox[1]
    draw.text((center[0] - w / 2, center[1] - h / 2), text, font=font, fill=fill)


def draw_rect_block(
    draw: ImageDraw.ImageDraw,
    box: tuple[int, int, int, int],
    title: str,
    lines: list[str],
    *,
    title_size: int = 34,
    body_size: int = 26,
) -> None:
    x1, y1, x2, y2 = box
    draw.rectangle(box, outline="black", fill="white", width=4)
    title_font = load_font(title_size, bold=True)
    body_font = load_font(body_size)

    wrapped_title = wrap_text(draw, title, title_font, x2 - x1 - 30)
    wrapped_body: list[str] = []
    for line in lines:
        wrapped_body.extend(wrap_text(draw, line, body_font, x2 - x1 - 34))

    title_h = title_size + 10
    body_h = body_size + 10
    total_h = len(wrapped_title) * title_h + 18 + len(wrapped_body) * body_h
    y = y1 + max(18, (y2 - y1 - total_h) // 2)
    cx = (x1 + x2) // 2

    for line in wrapped_title:
        draw_center_text(draw, (cx, y + title_size // 2), line, title_font)
        y += title_h
    y += 10
    for line in wrapped_body:
        draw_center_text(draw, (cx, y + body_size // 2), line, body_font)
        y += body_h


def draw_diamond(draw: ImageDraw.ImageDraw, center: tuple[int, int], w: int, h: int, text: str) -> None:
    cx, cy = center
    points = [(cx, cy - h // 2), (cx + w // 2, cy), (cx, cy + h // 2), (cx - w // 2, cy)]
    draw.polygon(points, outline="black", fill="white", width=4)
    draw_center_text(draw, center, text, load_font(30, bold=True))


def draw_line(draw: ImageDraw.ImageDraw, start: tuple[int, int], end: tuple[int, int], width: int = 4) -> None:
    draw.line([start, end], fill="black", width=width)


def draw_cardinality(draw: ImageDraw.ImageDraw, center: tuple[int, int], text: str) -> None:
    draw_center_text(draw, center, text, load_font(24, bold=True))


def generate_role_er(path: Path) -> None:
    img = Image.new("RGB", (2600, 1500), "white")
    draw = ImageDraw.Draw(img)

    draw_center_text(draw, (1300, 70), "系统角色需求ER图", load_font(56, bold=True))

    boundary = (220, 150, 2380, 1380)
    draw.rectangle(boundary, outline="black", width=5)
    draw.text((255, 180), "AI赋能高校教务系统", font=load_font(34, bold=True), fill="black")

    left_boxes = [
        ((340, 280, 760, 460), "学生", ["角色实体"]),
        ((340, 650, 760, 830), "教师", ["角色实体"]),
        ((340, 1020, 760, 1200), "管理员", ["角色实体"]),
    ]
    diamonds = [
        ((1030, 370), "对应"),
        ((1030, 740), "对应"),
        ((1030, 1110), "对应"),
    ]
    right_boxes = [
        ((1290, 220, 2200, 520), "学生侧需求实体", ["课程查询与办理", "AI客服咨询", "课程答疑支持"]),
        ((1290, 590, 2200, 890), "教师侧需求实体", ["课程资料上传与沉淀", "课程答疑支持", "智能教案生成与修改"]),
        ((1290, 960, 2200, 1260), "管理员侧需求实体", ["工作流编排与治理", "模型与知识库配置"]),
    ]

    for box, title, lines in left_boxes:
        draw_rect_block(draw, box, title, lines, title_size=36, body_size=24)
    for center, text in diamonds:
        draw_diamond(draw, center, 220, 140, text)
    for box, title, lines in right_boxes:
        draw_rect_block(draw, box, title, lines, title_size=34, body_size=28)

    rows = [
        ((760, 370), (920, 370), (1140, 370), (1290, 370)),
        ((760, 740), (920, 740), (1140, 740), (1290, 740)),
        ((760, 1110), (920, 1110), (1140, 1110), (1290, 1110)),
    ]
    for left_end, left_to_diamond, diamond_to_right, right_start in rows:
        draw_line(draw, left_end, left_to_diamond)
        draw_line(draw, diamond_to_right, right_start)

    draw_cardinality(draw, (840, 338), "1")
    draw_cardinality(draw, (1208, 338), "N")
    draw_cardinality(draw, (840, 708), "1")
    draw_cardinality(draw, (1208, 708), "N")
    draw_cardinality(draw, (840, 1078), "1")
    draw_cardinality(draw, (1208, 1078), "N")

    img.save(path)


def replace_image_and_caption(doc_path: Path, out_path: Path, image_path: Path) -> None:
    doc = Document(doc_path)
    caption_idx = next(i for i, p in enumerate(doc.paragraphs) if "图二-1 系统角色需求用例图" in p.text or "图二-1 系统角色需求ER图" in p.text)
    image_para = doc.paragraphs[caption_idx - 1]
    caption_para = doc.paragraphs[caption_idx]

    image_para.clear()
    image_para.add_run().add_picture(str(image_path), width=Cm(16))

    caption_para.clear()
    caption_para.add_run("图二-1 系统角色需求ER图")

    if out_path.exists():
        out_path.unlink()
    doc.save(out_path)


def main() -> None:
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    desktop = Path.home() / "Desktop"
    input_doc = next(desktop.glob("2405273202*chapter2-updated-v12.docx"))
    output_doc = input_doc.with_name(input_doc.stem.replace("-v12", "-v13") + input_doc.suffix)
    image_path = ASSET_DIR / "fig_2_1_role_demand_er_v1.png"

    generate_role_er(image_path)
    replace_image_and_caption(input_doc, output_doc, image_path)

    print(output_doc)
    print(image_path)


if __name__ == "__main__":
    main()
