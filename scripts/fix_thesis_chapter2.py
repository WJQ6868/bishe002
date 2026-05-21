from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.shared import Cm
from docx.text.paragraph import Paragraph
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
ASSET_DIR = ROOT / "generated" / "thesis_assets"
FONT_REGULAR = Path(r"C:\Windows\Fonts\msyh.ttc")
FONT_BOLD = Path(r"C:\Windows\Fonts\simhei.ttf")


def find_target_doc() -> Path:
    desktop = Path.home() / "Desktop"
    candidates = [
        p
        for p in desktop.glob("2405273202*567*.docx")
        if "backup" not in p.name.lower()
    ]
    if not candidates:
        raise FileNotFoundError("Target thesis document was not found on Desktop.")
    candidates.sort(key=lambda p: p.stat().st_mtime, reverse=True)
    return candidates[0]


def ensure_backup(doc_path: Path) -> Path:
    backup = doc_path.with_name(f"{doc_path.stem}-backup-before-chapter2-rewrite{doc_path.suffix}")
    if not backup.exists():
        shutil.copy2(doc_path, backup)
    return backup


def load_font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont:
    font_path = FONT_BOLD if bold and FONT_BOLD.exists() else FONT_REGULAR
    return ImageFont.truetype(str(font_path), size=size)


def wrap_text(draw: ImageDraw.ImageDraw, text: str, font: ImageFont.FreeTypeFont, max_width: int) -> list[str]:
    lines: list[str] = []
    current = ""
    for ch in text:
        candidate = current + ch
        if draw.textlength(candidate, font=font) <= max_width:
            current = candidate
        else:
            if current:
                lines.append(current)
            current = ch
    if current:
        lines.append(current)
    return lines or [text]


def draw_centered_text(draw: ImageDraw.ImageDraw, xy: tuple[int, int], text: str, font, fill=(44, 62, 80)) -> None:
    bbox = draw.textbbox((0, 0), text, font=font)
    width = bbox[2] - bbox[0]
    height = bbox[3] - bbox[1]
    draw.text((xy[0] - width / 2, xy[1] - height / 2), text, font=font, fill=fill)


def draw_arrow(draw: ImageDraw.ImageDraw, start: tuple[int, int], end: tuple[int, int], color=(52, 73, 94), width: int = 6) -> None:
    import math

    draw.line([start, end], fill=color, width=width)
    angle = math.atan2(end[1] - start[1], end[0] - start[0])
    arrow_len = 18
    angle1 = angle - math.pi / 8
    angle2 = angle + math.pi / 8
    p1 = (end[0] - arrow_len * math.cos(angle1), end[1] - arrow_len * math.sin(angle1))
    p2 = (end[0] - arrow_len * math.cos(angle2), end[1] - arrow_len * math.sin(angle2))
    draw.polygon([end, p1, p2], fill=color)


def draw_round_box(
    draw: ImageDraw.ImageDraw,
    box: tuple[int, int, int, int],
    title: str,
    lines: list[str],
) -> None:
    x1, y1, x2, y2 = box
    draw.rounded_rectangle(box, radius=24, fill=(255, 255, 255), outline=(52, 73, 94), width=4)
    title_font = load_font(44, bold=True)
    body_font = load_font(28)
    draw.text((x1 + 18, y1 + 16), title, font=title_font, fill=(52, 73, 94))
    y = y1 + 82
    max_width = x2 - x1 - 36
    for line in lines:
        for item in wrap_text(draw, line, body_font, max_width):
            draw.text((x1 + 18, y), item, font=body_font, fill=(33, 37, 41))
            y += 38


def generate_ai_closed_loop_only(path: Path) -> None:
    img = Image.new("RGB", (2400, 760), (255, 255, 255))
    draw = ImageDraw.Draw(img)
    draw_centered_text(draw, (1200, 58), "AI业务闭环逻辑", load_font(50, bold=True))

    boxes = [
        ((80, 180, 560, 520), "配置端", ["管理员维护模型接入方式、知识资料与场景参数"]),
        ((660, 180, 1140, 520), "业务端", ["学生咨询、课程答疑与教案生成等入口调用AI能力"]),
        ((1240, 180, 1720, 520), "数据沉淀", ["保存问答记录、课程资料、任务结果与处理状态"]),
        ((1820, 180, 2300, 520), "持续优化", ["依据使用反馈补充知识内容并调整模型与场景配置"]),
    ]
    centers = []
    for box, title, lines in boxes:
        draw_round_box(draw, box, title, lines)
        x1, y1, x2, y2 = box
        centers.append(((x1 + x2) // 2, (y1 + y2) // 2))

    draw_arrow(draw, (560, centers[0][1]), (660, centers[1][1]))
    draw_arrow(draw, (1140, centers[1][1]), (1240, centers[2][1]))
    draw_arrow(draw, (1720, centers[2][1]), (1820, centers[3][1]))

    loop_points = [(2060, 520), (2060, 650), (320, 650), (320, 540)]
    for start, end in zip(loop_points, loop_points[1:]):
        draw_arrow(draw, start, end)
    draw_centered_text(draw, (1190, 610), "运行结果进入下一轮知识维护与策略调整", load_font(28))

    img.save(path)


def generate_ai_layered_only(path: Path) -> None:
    img = Image.new("RGB", (2400, 1100), (255, 255, 255))
    draw = ImageDraw.Draw(img)
    draw_centered_text(draw, (1200, 60), "AI分层设计逻辑", load_font(50, bold=True))

    boxes = [
        ((160, 150, 2240, 330), "第一层：AI配置编排层", ["面向管理员，负责模型接入、场景工作流、参数策略与能力开关的统一治理。"]),
        ((160, 390, 2240, 570), "第二层：知识组织层", ["面向制度文档、课程资料与教学材料，完成上传、整理、归类和持续更新，使后续调用有稳定知识来源。"]),
        ((160, 630, 2240, 810), "第三层：推理执行层", ["围绕用户问题或生成任务完成请求理解、知识匹配、上下文组织和结果生成，是AI能力真正执行的核心层。"]),
        ((160, 870, 2240, 1050), "第四层：业务接入层", ["将统一能力开放给AI客服、AI课程助手和智能教案等页面入口，让不同角色按场景直接使用。"]),
    ]
    for box, title, lines in boxes:
        draw_round_box(draw, box, title, lines)

    draw_arrow(draw, (1200, 330), (1200, 390))
    draw_arrow(draw, (1200, 570), (1200, 630))
    draw_arrow(draw, (1200, 810), (1200, 870))
    img.save(path)


def para_index(doc: Document, text: str) -> int:
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip() == text:
            return i
    raise ValueError(f"Paragraph not found: {text}")


def find_para(doc: Document, text: str) -> Paragraph:
    return doc.paragraphs[para_index(doc, text)]


def set_paragraph_text(paragraph: Paragraph, text: str) -> None:
    paragraph.clear()
    paragraph.add_run(text)


def replace_following_paragraphs(doc: Document, heading: str, texts: list[str]) -> None:
    start = para_index(doc, heading)
    for offset, text in enumerate(texts, start=1):
        set_paragraph_text(doc.paragraphs[start + offset], text)


def find_table_by_header(doc: Document, first_cell_text: str):
    for table in doc.tables:
        if table.rows and table.rows[0].cells and table.rows[0].cells[0].text.strip() == first_cell_text:
            return table
    raise ValueError(f"Table not found: {first_cell_text}")


def move_body_block_after(doc: Document, start_para: Paragraph, end_para: Paragraph, after_para: Paragraph) -> None:
    body = doc._element.body
    children = list(body)
    start_idx = children.index(start_para._element)
    end_idx = children.index(end_para._element)
    insert_idx = children.index(after_para._element) + 1
    block = children[start_idx:end_idx]
    for element in block:
        body.remove(element)
    if insert_idx > start_idx:
        insert_idx -= len(block)
    for offset, element in enumerate(block):
        body.insert(insert_idx + offset, element)


def insert_paragraph_after(paragraph: Paragraph, text: str = "", style_name: str | None = None) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style_name:
        new_para.style = style_name
    if text:
        new_para.add_run(text)
    return new_para


def update_outline_paragraph(doc: Document) -> None:
    for para in doc.paragraphs:
        if "第二章从业务场景出发" in para.text:
            new_text = para.text.replace(
                "角色需求、功能需求、非功能需求、系统模块划分、总体架构、数据库总体规划以及AI闭环设计",
                "角色需求、功能需求、非功能需求、总体架构、系统模块划分、数据库总体规划以及AI闭环设计",
            )
            set_paragraph_text(para, new_text)
            return


def update_tables(doc: Document) -> None:
    role_table = find_table_by_header(doc, "角色")
    role_table.rows[2].cells[1].text = "课程管理、作业管理、教学组织"
    role_table.rows[2].cells[2].text = "课程资料上传、课程助手创建、智能教案生成"
    role_table.rows[3].cells[2].text = "模型接口管理、知识库管理、工作流编排、客服参数维护"

    module_table = find_table_by_header(doc, "模块")
    if module_table.rows[0].cells[1].text.strip() == "主要内容":
        module_table.rows[2].cells[1].text = "模型接口、知识库、工作流、客服参数"


def rewrite_chapter_two(doc: Document) -> None:
    replace_following_paragraphs(
        doc,
        "业务场景与角色需求分析",
        [
            "第二章的需求分析并不是先假定一个理想化场景，再去为AI寻找落点，而是回到高校教务系统本身来讨论问题。学生、教师和管理员共用同一套平台，但他们面对的业务压力、操作习惯和信息诉求并不一致。也正因为如此，AI能力不能脱离课程、请假、办事、成绩等既有业务独立存在，而应嵌入这些高频环节，承担解释、辅助和减负的作用。",
            "对学生而言，最常见的问题往往不是“系统里有没有这个功能”，而是“这个功能该怎么用、去哪里找、材料要准备什么”。传统教务系统能够提供规范入口，却未必能够及时说明流程差异和细节条件。因此，学生侧更需要一种边查询边理解的支持方式，在查看课程、办理请假、进入办事大厅或学习课程内容时，可以直接获得自然语言说明，减少来回跳转和重复咨询。",
            "教师侧的需求更偏向教学组织。教师既要完成课程管理、成绩处理、作业布置等确定性事务，又要承担备课、资料整理、课程答疑这些重复性较高的工作。现有教务系统在流程管理方面较成熟，但在教学内容辅助方面仍有提升空间，所以教师更希望系统能够围绕具体课程沉淀资料、支持课程问答，并在备课阶段提供可继续修改的生成结果。",
            "管理员的关注点又明显不同。管理员并不是主要提问者，但必须保证AI能力可以被稳定、可控地使用。这就带来了几项很现实的需求：模型接入方式要清楚，知识资料的归属要明确，不同场景要有对应的工作流，配置变更后前台页面还要能够及时生效。换句话说，管理员面对的是治理需求，而不是单纯的使用需求。",
            "如果把三类角色放到同一条业务线上看，系统至少要支撑三种典型场景。第一类是学生咨询场景，围绕选课、请假、办事和学习问题提供即时解释；第二类是课程答疑场景，由教师整理课程资料、学生围绕课程继续提问；第三类是教学生成场景，教师依据已有课程资源形成教案草稿并再加工。这些场景彼此并不割裂，它们都依赖课程数据、角色权限和知识资料的协同。",
            "因此，本系统的需求分析不能简单理解为给传统教务系统外接一个聊天窗口。传统教务系统负责规则执行、信息展示和流程办理，AI部分更像嵌入式能力层，负责解释制度、理解问题、组织资料和辅助生成。两部分分工清楚，系统才不会出现功能重复、角色边界模糊或结果难以落地的问题。",
            "从需求落点上看，学生侧强调可问、可懂、可继续办理；教师侧强调可沉淀、可复用、可追踪；管理员侧强调可配置、可控制、可维护。后续的功能分析、架构设计和数据库规划，实际上都是围绕这三类需求逐步展开的。",
        ],
    )

    replace_following_paragraphs(
        doc,
        "功能需求分析",
        [
            "在功能层面，系统并没有把AI能力单独剥离出来，而是以现有教务业务为底座，在其上增加面向咨询、学习和教学准备的增强能力。这样设计的原因很直接：学生和教师提出的问题，本来就发生在课程、请假、办事和教学等具体场景里。如果AI脱离这些业务上下文，即便能回答问题，实际使用价值也会明显下降。",
            "因此，基础教务功能首先需要保持完整和稳定。学生要能够完成课程查询、请假申请、成绩与证书查看、办事事项办理；教师要能够完成课程管理、作业与成绩处理、审批等教学组织工作；管理员要能够维护用户、流程和系统配置。这部分功能决定了系统的业务边界，也为AI判断面向谁服务、在什么场景下服务提供了前提。",
            "在此基础上，AI客服的功能需求主要体现为规则解释与入口引导。它要承接校园常见问题、办事咨询和流程说明，在页面中给出较为及时的反馈，并在知识不足时明确提示用户继续查看哪类业务入口或联系哪一类管理人员。这里的重点不是让AI代替业务办理，而是帮助用户更快理解业务。",
            "AI课程助手的功能需求更贴近课程场景。系统需要支持教师围绕具体课程整理资料、形成课程专属问答入口，学生再根据所选课程获取学习支持。这样一来，问答内容就不再停留在通用校园咨询，而是能够落到某门课程的知识点、资料范围和教学要求上。对教师而言，这项功能的核心价值在于把分散资料逐步转化为可持续复用的课程服务能力。",
            "智能教案模块则属于更典型的任务型功能。教师在备课时通常不会满足于一次性生成一段文本，而是希望系统围绕课程主题、资料基础和教学目标持续给出可修改、可保存、可导出的结果。所以这一模块的需求重点应放在过程完整性上，包括任务发起、生成反馈、结果编辑、历史保留和再次使用，而不是单独强调生成速度。",
            "从整体上看，AI增强功能可以归纳为三类需求：一类是配置类需求，确保管理员能够为不同场景准备合适的模型和知识资源；一类是交互类需求，确保学生和教师能在合适页面发起提问或生成请求，并清楚感知系统状态；另一类是沉淀类需求，确保文档、问答结果和教案内容能够被保存下来，服务下一轮课程答疑、教学准备和系统优化。这样划分之后，AI客服、AI课程助手和智能教案虽然面向对象不同，但在系统中的定位就比较清楚了。",
        ],
    )

    replace_following_paragraphs(
        doc,
        "非功能需求分析",
        [
            "除功能本身外，系统还需要满足较为明确的非功能需求。首先是可用性需求。高校教务业务往往具有时间敏感和流程敏感的特点，用户常常是在选课、请假、成绩查询或备课过程中临时进入系统，如果页面入口不清晰、反馈不及时，AI能力即便已经接入，也很难形成稳定的使用习惯。因此，系统需要在不同角色下提供清楚的入口组织和易于理解的操作反馈。",
            "其次是响应连续性需求。问答和内容生成与传统表单提交不同，用户更关心系统是否正在处理。为此，系统需要让返回过程具备可感知的连续反馈，而不是长时间无响应。论文中采用的服务器发送事件（Server-Sent Events，SSE）机制，本质上对应的就是这一类需求：让用户在等待结果时持续看到系统状态变化，从而降低不确定感。",
            "再次是可靠性与可追踪性需求。无论是课程资料上传、课程问答，还是教案生成，用户都希望知道操作是否成功、当前处于什么状态、出现问题后能否定位原因。尤其是教师侧的课程助手和智能教案，一旦缺少过程记录，后续修改、复查和复用都会受到影响。因此，系统不仅要能生成，还要能说明过程、能保留结果。",
            "安全与权限控制同样不可忽视。教务系统包含用户身份、课程信息、申请记录和教学资料等数据，不同角色的可见范围和可操作范围必须区分清楚。学生更偏向使用结果，教师需要管理课程资料，管理员负责配置和治理，这种权限差异决定了系统必须在访问入口、配置操作和数据归属上保持边界清晰，避免出现越权访问或资源混用。",
            "最后是部署与维护需求。毕业设计的展示并不只发生在本地开发环境中，还要考虑教师评审、远程访问和后续演示的实际需要。因此，系统应同时支持本地联调与公网访问场景。当前项目除支持本地启动外，还通过部署在阿里云云服务器 Ubuntu 系统上的站点对外提供服务，并可通过公网域名wangjiaqi.me直接访问。这类需求虽然不直接体现为页面功能，却决定了系统能否稳定演示、持续维护和方便扩展。",
        ],
    )

    replace_following_paragraphs(
        doc,
        "数据库总体规划",
        [
            "数据库总体规划是总体设计中的基础工作，它不只是保存页面结果，还承担着组织系统运行逻辑的作用。结合本项目的业务特点，数据库可以按数据职责分为五个部分：第一类是用户、课程、请假、成绩、办事等基础教务数据，用来支撑系统原有业务流程；第二类是模型、知识库与工作流等AI资源配置数据，用来说明AI能力由谁维护、面向什么场景开放；第三类是课程资料、制度文档及其整理结果，用来形成可持续调用的知识来源；第四类是教案生成等任务过程数据，用来记录任务发起、处理中和完成后的状态变化；第五类是问答记录、调用状态与运行痕迹数据，用来支撑后续复查、统计与优化[8]。",
            "从组织方式上看，这些数据并不是彼此孤立保存，而是围绕用户、课程、知识资料和AI场景逐层关联。基础教务数据负责提供角色身份和课程边界，AI配置数据负责确定调用关系，知识数据负责提供内容依据，任务与日志数据负责保留运行结果。这样的规划能够较清楚地回答三个问题：AI能力由谁配置、依赖哪些资料、最终结果保存在哪里，也使系统在后续扩展课程答疑、教学生成或使用统计时具备较好的结构基础。",
        ],
    )

    replace_following_paragraphs(
        doc,
        "AI闭环设计",
        [
            "结合本项目的业务组织方式，可以把AI功能主线概括为“配置准备、场景调用、结果沉淀、持续优化”四个连续环节。首先由管理员完成模型、知识资料和场景参数的准备，使系统具备可调用的AI底座；随后由学生和教师在具体页面中发起咨询、课程答疑或教案生成请求；系统运行后会把问答记录、课程资料变化、任务状态和生成结果保留下来；这些沉淀的数据又会反过来支持知识补充、配置调整和场景优化。这样形成的并不是一次性的调用链，而是一条能够持续迭代的业务闭环。",
            "为了让这一闭环能够稳定落到系统结构中，本文进一步将AI部分概括为四层设计：AI配置编排层负责模型接入和场景组织，知识组织层负责整理制度文档与课程资料，推理执行层负责完成问题理解、知识匹配和内容生成，业务接入层负责把这些能力开放给AI客服、AI课程助手和智能教案等具体入口[9][10][11][12]。这种写法强调的是设计层次，而不是把实现细节直接搬到需求分析中，也更符合本章从总体设计出发的叙述逻辑。",
        ],
    )


def split_ai_figure(doc: Document, closed_loop_path: Path, layered_path: Path) -> None:
    caption_para = find_para(doc, "图2-2 AI业务闭环与分层实现逻辑")
    caption_idx = para_index(doc, "图2-2 AI业务闭环与分层实现逻辑")
    image_para = doc.paragraphs[caption_idx - 1]
    image_para.clear()
    image_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    image_para.style = "样式1"
    image_para.add_run().add_picture(str(closed_loop_path), width=Cm(16.0))

    set_paragraph_text(caption_para, "图2-2 AI业务闭环逻辑")
    caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_para.style = "图片标题"

    second_image = insert_paragraph_after(caption_para, style_name="样式1")
    second_image.alignment = WD_ALIGN_PARAGRAPH.CENTER
    second_image.add_run().add_picture(str(layered_path), width=Cm(16.0))

    second_caption = insert_paragraph_after(second_image, "图2-3 AI分层设计逻辑", style_name="图片标题")
    second_caption.alignment = WD_ALIGN_PARAGRAPH.CENTER


def main() -> None:
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    doc_path = find_target_doc()
    ensure_backup(doc_path)

    fig_2_2 = ASSET_DIR / "fig_2_2_ai_closed_loop_only.png"
    fig_2_3 = ASSET_DIR / "fig_2_3_ai_layered_design.png"
    generate_ai_closed_loop_only(fig_2_2)
    generate_ai_layered_only(fig_2_3)

    doc = Document(doc_path)
    update_outline_paragraph(doc)
    rewrite_chapter_two(doc)
    update_tables(doc)

    move_body_block_after(
        doc,
        find_para(doc, "系统模块划分"),
        find_para(doc, "系统总体架构设计"),
        find_para(doc, "图2-1 系统总体架构图"),
    )

    split_ai_figure(doc, fig_2_2, fig_2_3)
    doc.save(doc_path)

    print(f"Updated: {doc_path}")
    print(f"Backup: {doc_path.with_name(f'{doc_path.stem}-backup-before-chapter2-rewrite{doc_path.suffix}')}")


if __name__ == "__main__":
    main()
